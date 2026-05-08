"""Cover-letter export helpers (CL-9).

v2.1 of the cover-letter feature stores `content` as a Tiptap JSON
document (`{type: "doc", content: [paragraph nodes...]}`), so the
legacy plain-text exporters that lived in this module are gone — the
new world is JSON-in, bytes-out.

Two pure-Python helpers live here:

* ``render_tiptap_to_html`` — walks the Tiptap doc and emits a
  styled, letter-aesthetic HTML document. The output is fed to the
  Playwright/Chromium pipeline (``app.services.pdf_service``) for
  PDF rendering — no JS runtime required.

* ``render_tiptap_to_docx`` — uses ``python-docx`` to materialise
  the same doc as a Word file. Bold/italic marks and ``hardBreak``
  nodes are preserved at the run level.

Other rich Tiptap node types (lists, headings, links, …) are out of
scope for the wizard — Variants are paragraph-only — so any unknown
node type is rendered as plain concatenated text via
``_collect_text``. That keeps the exporters forgiving in case a
future PRD widens the editor without an immediate exporter update.
"""
from __future__ import annotations

import html as html_lib
import io
from typing import Any, Optional


# ---------------------------------------------------------------------------
# Tiptap structural constants
# ---------------------------------------------------------------------------
# Names lifted directly from prosemirror-model node specs the Tiptap
# wizard emits — keeping them as constants avoids typo bugs and makes
# the walker self-documenting.

_NODE_DOC = "doc"
_NODE_PARAGRAPH = "paragraph"
_NODE_TEXT = "text"
_NODE_HARD_BREAK = "hardBreak"

_MARK_BOLD = "bold"
_MARK_ITALIC = "italic"


# ---------------------------------------------------------------------------
# Mark detection
# ---------------------------------------------------------------------------


def _has_mark(node: dict, mark_type: str) -> bool:
    """Return True if a Tiptap text node carries the named mark.

    Tiptap stores marks as a list of ``{type, attrs}`` dicts on the
    node itself. We only care about the ``type`` for bold/italic — no
    attrs to inspect.
    """
    marks = node.get("marks") or []
    for mark in marks:
        if isinstance(mark, dict) and mark.get("type") == mark_type:
            return True
    return False


def _collect_text(node: Any) -> str:
    """Best-effort plain-text walker for nodes we don't render explicitly.

    Used only for unknown / unsupported node types so the export never
    silently drops user-typed content. Concatenates every text leaf
    found in the subtree, separated by nothing — Tiptap text runs that
    were adjacent in the source remain adjacent in the fallback.
    """
    if isinstance(node, dict):
        if node.get("type") == _NODE_TEXT:
            return str(node.get("text") or "")
        children = node.get("content") or []
        return "".join(_collect_text(child) for child in children)
    if isinstance(node, list):
        return "".join(_collect_text(child) for child in node)
    return ""


# ---------------------------------------------------------------------------
# Tiptap → HTML
# ---------------------------------------------------------------------------


def _render_paragraph_html(paragraph: dict) -> str:
    """Render one ``paragraph`` node as an HTML ``<p>...</p>`` block.

    Empty paragraphs (no children) are emitted as
    ``<p class="empty">&nbsp;</p>`` — the ``empty`` class is what
    shrinks the spacer (see the ``p.empty`` rule in
    ``render_tiptap_to_html``) so blank lines stay visible but tight,
    instead of inheriting the full body line-height + paragraph
    margin which produces ~46px gaps.
    """
    children = paragraph.get("content") or []
    if not children:
        return '<p class="empty">&nbsp;</p>'

    parts: list[str] = []
    for child in children:
        if not isinstance(child, dict):
            continue
        ctype = child.get("type")

        if ctype == _NODE_TEXT:
            text = html_lib.escape(str(child.get("text") or ""))
            if _has_mark(child, _MARK_BOLD):
                text = f"<strong>{text}</strong>"
            if _has_mark(child, _MARK_ITALIC):
                text = f"<em>{text}</em>"
            parts.append(text)
        elif ctype == _NODE_HARD_BREAK:
            parts.append("<br>")
        else:
            # Unknown inline node — fall back to escaped plain text so
            # nothing is silently dropped if the wizard adds new types.
            parts.append(html_lib.escape(_collect_text(child)))

    return f"<p>{''.join(parts)}</p>"


def render_tiptap_to_html(doc: dict, *, title: str = "Cover Letter") -> str:
    """Render a Tiptap ``doc`` to a complete styled HTML document.

    Styling target: a minimalist letter-style page.
      * 11pt body, 1.4 line-height — comfortable letter density
        without the spread-out feel a 1.5 setting gave the closing
        contact block.
      * A neutral system font stack so the PDF doesn't depend on any
        bundled font being available to Chromium (cover letters don't
        need the resume's custom-font story).

    Page margins are intentionally NOT set here — they are enforced
    solely by ``app.services.pdf_service.render_pdf``'s ``margin_mm``
    parameter, which Chromium applies via ``page.pdf({margin: ...})``.
    Adding any wrapper padding or ``@page`` margin in this HTML would
    stack on top of that and produce excessive whitespace in the PDF.
    """
    children = doc.get("content") or []

    body_blocks: list[str] = []
    for node in children:
        if not isinstance(node, dict):
            continue
        ntype = node.get("type")
        if ntype == _NODE_PARAGRAPH:
            body_blocks.append(_render_paragraph_html(node))
        else:
            # Unknown top-level node — render as a paragraph of its
            # collected text so we never drop content silently.
            text = html_lib.escape(_collect_text(node))
            body_blocks.append(f"<p>{text or '&nbsp;'}</p>")

    body_html = "\n".join(body_blocks) or "<p>&nbsp;</p>"
    safe_title = html_lib.escape(title)

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{safe_title}</title>
<style>
* {{ box-sizing: border-box; }}
html, body {{
  margin: 0;
  padding: 0;
  font-family: Georgia, "Times New Roman", Times, serif;
  font-size: 11pt;
  line-height: 1.4;
  color: #111;
  background: #fff;
}}
p {{
  margin: 0 0 6px 0;
  white-space: pre-wrap;
}}
p.empty {{
  margin: 0;
  line-height: 1;
  height: 0.6em;
}}
</style>
</head>
<body>
{body_html}
</body>
</html>"""


# ---------------------------------------------------------------------------
# Tiptap → DOCX
# ---------------------------------------------------------------------------


def _add_paragraph_to_docx(doc, paragraph: dict) -> None:  # type: ignore[no-untyped-def]
    """Append one Tiptap ``paragraph`` to a ``python-docx`` Document.

    Empty paragraphs (no children) become a blank ``add_paragraph()``
    so the visual line-break between letter sections is preserved.

    Bold / italic marks are mapped onto run-level attributes; a
    ``hardBreak`` node within a paragraph emits ``run.add_break()``
    and starts a fresh run (so subsequent marks don't bleed across
    the break).
    """
    children = paragraph.get("content") or []
    para = doc.add_paragraph()
    if not children:
        return

    for child in children:
        if not isinstance(child, dict):
            continue
        ctype = child.get("type")

        if ctype == _NODE_TEXT:
            run = para.add_run(str(child.get("text") or ""))
            if _has_mark(child, _MARK_BOLD):
                run.bold = True
            if _has_mark(child, _MARK_ITALIC):
                run.italic = True
        elif ctype == _NODE_HARD_BREAK:
            # Add the line-break to whichever run is current; if there
            # isn't one yet, create an empty placeholder run for it.
            if para.runs:
                para.runs[-1].add_break()
            else:
                para.add_run().add_break()
        else:
            # Unknown inline node — preserve user text via the fallback
            # walker. No marks (we don't know what they'd mean here).
            text = _collect_text(child)
            if text:
                para.add_run(text)


def render_tiptap_to_docx(doc: dict, *, title: Optional[str] = None) -> bytes:
    """Render a Tiptap ``doc`` to a DOCX byte stream.

    ``title`` is intentionally unused at the document body level — the
    cover-letter file already contains the salutation/closing the user
    wrote, so injecting a title heading would duplicate content. The
    parameter is kept for symmetry with the HTML renderer and a future
    "show name as heading" toggle.
    """
    # Local import — keeps this module importable in environments where
    # python-docx isn't on the path (mirrors the pdf_service pattern).
    from docx import Document
    from docx.shared import Inches, Pt

    document = Document()

    # 1in uniform margin — matches the HTML/PDF aesthetic.
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Default body font — set on the Normal style so all runs inherit
    # it unless the Tiptap node carries an explicit override.
    style = document.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    children = doc.get("content") or []
    for node in children:
        if not isinstance(node, dict):
            continue
        ntype = node.get("type")
        if ntype == _NODE_PARAGRAPH:
            _add_paragraph_to_docx(document, node)
        else:
            # Unknown top-level node — drop its text into a paragraph
            # so nothing is silently lost.
            text = _collect_text(node)
            if text:
                document.add_paragraph(text)

    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()
