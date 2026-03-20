"""
DOCX export utility for resumes.

Takes a parsed_data dict (the ResumeData structure stored in the DB) and
returns a well-formatted DOCX document as bytes.
"""
from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def generate_docx(parsed_data: dict[str, Any]) -> bytes:
    """
    Convert a parsed_data dict into a DOCX file and return it as bytes.

    Sections are rendered in the order defined by parsed_data keys, which
    mirrors the sections_order stored on the Resume model:
      personal_info, summary, experience, education, skills, languages,
      certifications.
    """
    doc = Document()
    _configure_margins(doc)

    personal_info: dict[str, Any] = parsed_data.get("personal_info") or {}
    _render_personal_info(doc, personal_info)

    summary: str | None = parsed_data.get("summary")
    if summary and summary.strip():
        _render_section_heading(doc, "Summary")
        doc.add_paragraph(summary.strip())

    experience: list[dict] = parsed_data.get("experience") or []
    if experience:
        _render_section_heading(doc, "Experience")
        for item in experience:
            _render_experience_item(doc, item)

    education: list[dict] = parsed_data.get("education") or []
    if education:
        _render_section_heading(doc, "Education")
        for item in education:
            _render_education_item(doc, item)

    skills: list[str] = parsed_data.get("skills") or []
    if skills:
        _render_section_heading(doc, "Skills")
        doc.add_paragraph(", ".join(s for s in skills if s))

    languages: list[str] = parsed_data.get("languages") or []
    if languages:
        _render_section_heading(doc, "Languages")
        doc.add_paragraph(", ".join(l for l in languages if l))

    certifications: list[dict] = parsed_data.get("certifications") or []
    if certifications:
        _render_section_heading(doc, "Certifications")
        for cert in certifications:
            _render_certification_item(doc, cert)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------

def _render_personal_info(doc: Document, info: dict[str, Any]) -> None:
    name: str = (info.get("name") or "").strip()
    if name:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(name)
        run.bold = True
        run.font.size = Pt(20)

    contact_parts: list[str] = []
    for field in ("email", "phone", "location", "linkedin", "website"):
        value = (info.get(field) or "").strip()
        if value:
            contact_parts.append(value)

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(10)


def _render_section_heading(doc: Document, heading: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(heading.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Draw a bottom border on the paragraph to act as a divider
    _add_paragraph_bottom_border(para)


def _render_experience_item(doc: Document, item: dict[str, Any]) -> None:
    title: str = (item.get("title") or "").strip()
    company: str = (item.get("company") or "").strip()
    start_date: str = (item.get("start_date") or "").strip()
    end_date: str = ("Present" if item.get("current") else (item.get("end_date") or "")).strip()
    description: str = (item.get("description") or "").strip()
    achievements: list[str] = item.get("achievements") or []

    # Job title and company on one line
    heading_parts = []
    if title:
        heading_parts.append(title)
    if company:
        heading_parts.append(company)
    heading_text = " — ".join(heading_parts) if heading_parts else "Untitled Position"

    para = doc.add_paragraph()
    run = para.add_run(heading_text)
    run.bold = True
    run.font.size = Pt(11)

    # Dates on the same paragraph, right-aligned via a tab stop is complex —
    # append as a separate, lighter line instead.
    date_parts = [p for p in (start_date, end_date) if p]
    if date_parts:
        date_para = doc.add_paragraph(" – ".join(date_parts))
        for run in date_para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if description:
        doc.add_paragraph(description)

    for achievement in achievements:
        if achievement and achievement.strip():
            doc.add_paragraph(achievement.strip(), style="List Bullet")


def _render_education_item(doc: Document, item: dict[str, Any]) -> None:
    institution: str = (item.get("institution") or "").strip()
    degree: str = (item.get("degree") or "").strip()
    field: str = (item.get("field") or "").strip()
    start_date: str = (item.get("start_date") or "").strip()
    end_date: str = (item.get("end_date") or "").strip()
    gpa: str = (item.get("gpa") or "").strip()

    degree_parts = [p for p in (degree, field) if p]
    degree_text = ", ".join(degree_parts) if degree_parts else ""

    heading_parts = [p for p in (institution, degree_text) if p]
    heading_text = " — ".join(heading_parts) if heading_parts else "Untitled"

    para = doc.add_paragraph()
    run = para.add_run(heading_text)
    run.bold = True
    run.font.size = Pt(11)

    date_parts = [p for p in (start_date, end_date) if p]
    meta_parts = []
    if date_parts:
        meta_parts.append(" – ".join(date_parts))
    if gpa:
        meta_parts.append(f"GPA: {gpa}")

    if meta_parts:
        meta_para = doc.add_paragraph(" | ".join(meta_parts))
        for run in meta_para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _render_certification_item(doc: Document, item: dict[str, Any]) -> None:
    name: str = (item.get("name") or "").strip()
    issuer: str = (item.get("issuer") or "").strip()
    date: str = (item.get("date") or "").strip()

    parts = [p for p in (name, issuer, date) if p]
    if parts:
        doc.add_paragraph(" | ".join(parts), style="List Bullet")


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _configure_margins(doc: Document) -> None:
    """Set 1-inch margins on all sides for the default section."""
    from docx.shared import Inches
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def _add_paragraph_bottom_border(para) -> None:  # type: ignore[no-untyped-def]
    """Add a thin bottom border to a paragraph via raw OOXML manipulation."""
    from lxml import etree

    pPr = para._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")   # half-points: 4 = 0.5pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
