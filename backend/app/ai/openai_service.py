from __future__ import annotations

import json
import logging
import time
from typing import Any, Optional

from openai import AsyncOpenAI

from app.ai.interface import (
    AIServiceInterface,
    AIUsageInfo,
    CoverLetterVariant,
    GenerateCoverLetterResult,
    GenerateTailoredResumeResult,
    ParsedJobData,
    ParseJobResult,
    ParseResumeProfileResult,
)
from app.config import settings

logger = logging.getLogger(__name__)

_MODEL = "gpt-4o-mini"
_TIMEOUT = 30.0        # seconds — for simple calls (job parsing)
_PROFILE_PARSE_TIMEOUT = 60.0  # seconds — for resume → ProfileData parsing (PROFILE-2)
_TAILOR_TIMEOUT = 90.0  # seconds — for tailored resume generation (TAILOR-1)
_COVER_LETTER_TIMEOUT = 90.0  # seconds — for 3-variant CL generation (CL-1)
_TAILOR_MAX_KEYWORDS = 12  # cap for matched_keywords (UI side panel)
_TAILOR_PALETTE_SIZE = 8   # color_id cycles through 1..8

# CL length bands (PRD 3.5.4). The AI is instructed to land within ±15% of
# the band's midpoint. Post-validation only enforces non-empty content; word
# count drift is logged-but-not-rejected so a slightly over/under variant is
# still surfaced to the user rather than failing the whole batch.
_CL_LENGTH_BANDS: dict[str, tuple[int, int]] = {
    "short": (200, 300),
    "medium": (300, 400),
    "long": (400, 500),
}
_CL_VARIANT_COUNT = 3

# Hosts whose profile URLs always require a path component. A bare root URL
# like "https://www.linkedin.com" with no /in/<username> is not a profile and
# must be rejected from the parsed social_links.
_PLATFORM_ROOT_HOSTS: frozenset[str] = frozenset({
    "linkedin.com",
    "www.linkedin.com",
    "github.com",
    "www.github.com",
    "x.com",
    "www.x.com",
    "twitter.com",
    "www.twitter.com",
    "facebook.com",
    "www.facebook.com",
    "instagram.com",
    "www.instagram.com",
    "youtube.com",
    "www.youtube.com",
    "dribbble.com",
    "www.dribbble.com",
    "behance.net",
    "www.behance.net",
})

_JOB_PARSE_SYSTEM_PROMPT = """\
You are an expert recruiter and job description analyst. Your task is to extract structured information from job descriptions provided by the user.

The job description may be in any format: formal corporate postings, informal startup descriptions, LinkedIn posts, emails, bullet-point lists, or paragraphs. It may be written in any language.

Extract the following fields:
- job_title: The title of the position (e.g. "Senior Backend Engineer"). If unclear, make your best inference. If truly absent, return null.
- company: The name of the hiring company or organisation. If the posting is confidential or uses a recruiting agency without naming the client, return null.
- requirements: A flat list of individual skills, qualifications, tools, or experience requirements. Each item must be a concise string (e.g. "Python", "5+ years backend experience", "Bachelor's degree in Computer Science"). Extract all meaningful requirements. Return an empty list if none are found.
- location: City, country, "Remote", "Hybrid", or a combination (e.g. "Berlin, Germany (Hybrid)"). Return null if not mentioned.
- salary: The compensation as stated, including currency and period if present (e.g. "$90,000–$120,000/year", "€60k", "£400/day"). Return null if not mentioned.

Rules:
- Never hallucinate or invent information that is not present in the text.
- If a field cannot be confidently determined from the text, return null (or [] for requirements).
- Return ONLY valid JSON with exactly these five keys: job_title, company, requirements, location, salary.
- Do not include any explanation, markdown, or extra text outside the JSON object.

Example output:
{
  "job_title": "Backend Engineer",
  "company": "Acme Corp",
  "requirements": ["Python", "FastAPI", "PostgreSQL", "3+ years experience"],
  "location": "Remote",
  "salary": "$80,000–$100,000/year"
}
"""

_PROFILE_PARSE_SYSTEM_PROMPT = """\
You are an expert resume parser. Extract a complete, structured user profile from the resume text provided by the user.

The resume may be in any format and any language. Read it carefully and map every piece of information to the correct section of the profile schema below.

Return ONLY a valid JSON object with EXACTLY these top-level keys (no extras, no missing keys):

{
  "personal_info": {
    "full_name": string,
    "email": string,
    "phone": string,
    "location": string | null,
    "social_links": [
      {"platform": string, "url": string}
    ]
  },
  "summary": string,
  "employment": [
    {
      "role": string,
      "company": string,
      "location": string | null,
      "start_date": string,
      "end_date": string | null,
      "is_current": boolean,
      "details": [string]
    }
  ],
  "education": [
    {
      "degree": string,
      "institution": string,
      "field_of_study": string | null,
      "location": string | null,
      "start_date": string,
      "end_date": string | null,
      "is_current": boolean,
      "description": string | null
    }
  ],
  "skills": [
    {"name": string}
  ],
  "projects": [
    {
      "name": string,
      "description": string,
      "role": string | null,
      "start_date": string | null,
      "end_date": string | null,
      "is_current": boolean,
      "technologies": [string],
      "url": string | null,
      "repository_url": string | null,
      "details": [string]
    }
  ],
  "languages": [
    {"name": string}
  ],
  "certificates": [
    {
      "name": string,
      "issuer": string | null,
      "issue_date": string | null,
      "expiry_date": string | null,
      "credential_url": string | null
    }
  ],
  "achievements": [
    {
      "title": string,
      "description": string | null,
      "date": string | null,
      "issuer": string | null
    }
  ],
  "volunteer": [
    {
      "role": string,
      "organization": string,
      "location": string | null,
      "start_date": string,
      "end_date": string | null,
      "is_current": boolean,
      "details": [string]
    }
  ]
}

CRITICAL RULES:

1. NEVER invent or hallucinate jobs, education, skills, projects, certifications, achievements, languages, or any other data not literally present in the resume text. If a section is missing from the resume, return an empty array for it.

2. Do NOT generate or include any `id` field on any item — IDs are assigned by the frontend.

3. Date format: Use "YYYY-MM" (year-month) for all date fields (start_date, end_date, issue_date, expiry_date, date). If only a year is given, use "YYYY-01" for start dates and "YYYY-12" for end dates. If a date is genuinely missing, use null (or empty string only where the schema requires a string — see rules 6 and 8).

4. Current positions: If a role/study/project says "Present", "Current", "Now", "to date" or has no end date listed for an ongoing item, set is_current=true AND end_date=null.

5. employment.details and volunteer.details and project.details: Each is an array of bullet-point strings. Split the source on bullet markers (•, -, *, ▪) or line breaks. Keep each bullet as one separate string. Preserve original wording. If only a paragraph is given (no bullets), put it as a single-element array.

6. personal_info.email: REQUIRED string. If the resume contains no email, use empty string "" (NOT null). The user can fix this after upload.

7. personal_info.phone: REQUIRED string. If missing, use empty string "" (NOT null).

8. personal_info.full_name: REQUIRED string. Use the candidate's name as it appears in the resume (typically the heading). If genuinely absent, use empty string "".

9. employment[].start_date and education[].start_date and volunteer[].start_date are REQUIRED strings. If the resume omits a start date for an entry, use empty string "" — do NOT invent dates.

10. summary: A plain-text string. Extract any "Profile", "Summary", "Professional Summary", "About Me", or "Objective" section. If none exists, use empty string "".

11. social_links: Extract any LinkedIn, GitHub, personal website, X (Twitter), portfolio, blog, or other social URLs found anywhere in the resume. For the `platform` field, use one of these whitelisted values that best fits: "LinkedIn", "GitHub", "X", "Portfolio", "Website", "Blog", "Facebook", "Instagram", "YouTube", "Dribbble", "Behance", "Custom". Use "Custom" if no other option matches. The `url` field MUST be a real, complete URL extracted from the resume — typically containing "://" or a domain like "linkedin.com/in/xyz" or "github.com/xyz". If the resume only mentions a platform name (e.g. just the word "LinkedIn" or a logo) WITHOUT an accompanying URL or username, OMIT that social_link entry entirely. NEVER use the platform name as the URL value, and NEVER invent or guess URLs. If you only have a username (e.g. "@johndoe" on GitHub), construct the canonical URL ("https://github.com/johndoe"). If you have a bare domain ("portfolio.dev"), prepend "https://". The user message may include a section labeled "[DETECTED_URLS]" near the end — these are URLs extracted from the document's hyperlink annotations (often hidden under icons or short visible text like "LinkedIn"). Use these URLs to populate `social_links` whenever you can correlate a URL to its platform by domain (linkedin.com → LinkedIn, github.com → GitHub, twitter.com or x.com → X, etc.). Prefer FULL profile URLs from [DETECTED_URLS] over short visible text. NEVER output a bare root-domain URL like "https://linkedin.com" or "https://www.github.com" — those have no profile path and are useless. If the resume gives you only a root domain with no username/path, OMIT the social_link entry entirely.

12. skills: Each is just `{"name": "<skill>"}` — do NOT include category or level. Split comma-separated or pipe-separated skill lists into individual items.

13. languages: Each is just `{"name": "<language>"}` — extract the language name only (e.g. "English", "German"). Do NOT include proficiency levels in the name field.

14. projects: Include personal projects, open-source contributions, side projects, hackathon projects. `description` is a one-paragraph summary; `details` is the array of bullet points (achievements/features). If the source only has bullets, put them in `details` and use the first bullet (or a brief synthesis) as `description`.

15. achievements: Include awards, hackathon wins, published works, recognitions, scholarships. NOT job-related accomplishments — those go into employment[].details.

16. certificates: Include only items the resume labels as certifications, certificates, or licenses (e.g. "AWS Certified Solutions Architect"). Do NOT promote skills or courses to certificates.

17. Output ONLY valid JSON — no markdown fences, no commentary, no preamble, no trailing text.
"""


_TAILOR_RESUME_SYSTEM_PROMPT = """\
You are an expert resume tailor and ATS-optimization specialist. You will be given the user's full PROFILE, a parsed JOB, and resume PREFERENCES. Your mission is to maximise the resume's chance of passing ATS keyword filters and reaching a human recruiter, without fabricating verifiable facts about the candidate.

WHAT YOU MAY DO (ATS optimization):
- ADD skills, frameworks, tools, methodologies, and soft skills that appear in the JOB and are plausibly within the candidate's professional domain given their existing experience. Insert them into the `skills` section. Prefer skills adjacent to what the candidate already does; avoid additions that would be jarringly outside their stated background (e.g. don't add "Rust embedded systems" to a pure frontend developer).
- REPHRASE `summary`, employment `details`, project `description`/`details`, and volunteer `details` to surface job-relevant vocabulary, problems-solved framings, and outcomes the recruiter cares about. If the JOB describes specific problems the employer wants solved, recast the candidate's existing accomplishments to emphasise that they have solved analogous problems, using only what the profile supports.
- REORDER items in any list section so the most relevant entries appear first.
- WEAVE job keywords naturally into rewritten prose. Do not keyword-stuff.

WHAT YOU MUST NEVER DO (factual integrity):
- Do NOT invent or add new entries to: `employment`, `education`, `certificates`, `volunteer`, `projects`, `languages`, `achievements`. These represent verifiable facts and must remain exactly the set the user provided.
- Do NOT invent or modify quantitative metrics. If the original bullet says "improved performance", you may rewrite it as "optimised performance for high-throughput workloads" but NOT "improved performance by 40%". Numbers, percentages, headcounts, revenue figures, durations, and timeframes that are not already in the profile are off-limits.
- Do NOT modify identifying facts in any existing entry: company names, institution names, role titles, dates, locations, certificate issuers, language proficiency levels, project names. Preserve these verbatim.
- Do NOT change languages spoken or their proficiency levels.
- The PROFILE you receive will have contact fields (email, phone, location, links) stripped from `personal_info` (this is intentional: contact data is removed before being sent to the model for privacy). Set `tailored_data.personal_info` to whatever was provided (possibly just a name, or null). Do NOT invent emails, phone numbers, or links.
- Preserve all `id` fields verbatim if present.

OUTPUT: return ONLY a single valid JSON object with EXACTLY these four top-level keys (no markdown fences, no extra text):

{
  "tailored_data": {
    "personal_info": { /* copied verbatim from the (already-sanitised) profile */ },
    "summary": "<rewritten to emphasise alignment with the job>",
    "employment": [ /* same items, may be reordered, details rewritten; company/title/dates/location preserved verbatim */ ],
    "education":  [ /* same items, may be reordered; all fields preserved verbatim */ ],
    "skills":     [ /* existing skills plus added job-relevant skills; may be reordered */ ],
    "projects":   [ /* same items, may be reordered, description/details rewritten; name/dates preserved verbatim */ ],
    "languages":  [ /* preserved verbatim; may be reordered */ ],
    "certificates":[ /* preserved verbatim; may be reordered */ ],
    "achievements":[ /* preserved verbatim; may be reordered */ ],
    "volunteer":  [ /* same items, may be reordered, details rewritten; organisation/title/dates preserved verbatim */ ]
  },
  "matched_keywords": [
    {"term": "<concrete skill/tool/qualification from the JOB>", "color_id": <integer 1..8>}
  ],
  "applied_changes": {
    "<section_key>": ["<natural-language description of what changed>", "..."]
  },
  "match_score": <integer 0..100>
}

RULES PER FIELD:

1. tailored_data MUST contain ALL of these keys, even when empty: personal_info, summary, employment, education, skills, projects, languages, certificates, achievements, volunteer. Use empty arrays for empty sections. Use the same shape as ProfileData (PRD 6.4). Copy `personal_info` verbatim from the input (it has already been sanitised by the backend). Preserve all `id` fields verbatim.

2. preferences.sections_order is a hint about which sections the user prioritises in their resume layout. Use this when deciding where to invest the most tailoring effort: sections appearing earlier matter more. The keys in `tailored_data` itself are always the canonical names listed in rule 1, regardless of order.

3. Skills additions:
   - Only add skills that appear in the JOB requirements/description AND are plausibly within the candidate's professional domain.
   - Place added skills among existing skills in a logical order (not all clustered at the top), so the list reads as a coherent skill set rather than a keyword dump.
   - Use the same shape and naming convention as existing skill entries in the profile.
   - In `applied_changes.skills`, explicitly list which skills were added (e.g. "Added 'Docker', 'CI/CD', 'PostgreSQL' to align with job requirements.") separately from any reordering note.

4. matched_keywords:
   - Extract concrete, job-specific terms from `job.requirements` and `job.description`. Prefer nouns and tools (e.g. "FastAPI", "PostgreSQL", "Kubernetes", "team lead", "GDPR") over generic verbs or fluff (e.g. "strong communicator", "team player").
   - Cap the list at 12 keywords. Choose the most important ones if there are more.
   - color_id MUST be an integer in 1..8. Cycle 1..8 deterministically by index (1st keyword=1, 2nd=2, ..., 8th=8, 9th=1, ...). Stable across calls, never randomize.
   - Each keyword appears at most once.

5. applied_changes:
   - Object keyed by section name (one of: "summary", "employment", "education", "skills", "projects", "languages", "certificates", "achievements", "volunteer").
   - Include ONLY sections you actually modified. Omit unchanged sections entirely.
   - Each value is a non-empty array of short, specific descriptions (e.g. "Rephrased opening line to lead with distributed-systems experience.", "Reordered Python and FastAPI to the top of skills.", "Added 'Kubernetes' and 'Helm' to skills based on job requirements.").
   - Be specific: reference the actual change, not boilerplate like "improved section".

6. match_score:
   - Integer 0..100 reflecting how well the ORIGINAL profile matches the job BEFORE tailoring (so the baseline gap is visible to the user).
   - Calibrate honestly: do not inflate. 90+ means strong alignment on most requirements; 50 to 70 partial alignment; below 40 significant gaps.

7. Language: Mirror the language of the profile and job content. If they differ, prefer the language of the job description for the tailored output (this is what the recruiter will read).

8. Punctuation style: Do NOT use em-dashes (the long dash character) or en-dashes anywhere in the rewritten prose (summary, employment/project/volunteer details, project descriptions, and applied_changes descriptions). Em-dashes are a recognizable AI-writing signature and must be avoided. Use commas, colons, periods, or parentheses instead. This applies to salary/range style too: prefer "to" or "from X to Y" over a dash between numbers.

Return ONLY the JSON object. No commentary, no preamble.
"""


_COVER_LETTER_SYSTEM_PROMPT = """\
You are an expert career coach and professional copywriter specialising in cover letters. You will be given a SOURCE (the applicant's profile or a tailored resume, same JSON shape either way), a JOB (parsed job posting), and PREFERENCES (style, tone, length, optional additional_context).

Your task: produce EXACTLY 3 cover letter variants in a SINGLE response. Each variant is a complete, ready-to-send letter (greeting, body, closing, signature line).

ABSOLUTE SOURCE-OF-TRUTH RULE (read carefully):
- The SOURCE is the ONLY allowed source of facts about the applicant. Do NOT invent companies, roles, dates, education, projects, certifications, achievements, or specific accomplishments that are not already present in the SOURCE.
- You MAY rephrase, summarise, and re-emphasise existing facts to better match the job. You MAY draw motivation/intent from `additional_context` if provided.
- The SOURCE is labelled with `source_type` ("tailored_resume" or "profile") for context only: DO NOT change your behaviour based on this field. The factual base is identical in either case.
- The SOURCE's personal_info has contact details (email, phone, location, links) removed for privacy; only the name may be present. Sign the letter using the applicant's name when available; never invent contact details.

VARIANTS MUST BE LEGIBLY DIFFERENT. Past attempts produced near-duplicates; that is unacceptable. Use these mandatory differentiators:
- Variant 1 (ACHIEVEMENT-FIRST): open with the applicant's strongest, most job-relevant accomplishment (a concrete result with numbers if available). Body emphasises proof of impact. Tight, evidence-driven structure.
- Variant 2 (MOTIVATION-FIRST): open with why the applicant is drawn to THIS role or problem space. Body bridges personal motivation to relevant experience. Narrative, slightly more personal arc.
- Variant 3 (COMPANY-ALIGNMENT-FIRST): open by referencing something specific to the company or role from the JOB (mission, product area, requirement, stated challenge). Body shows how the applicant's background slots into that need. Most "tailored to this employer" feel.

All three variants share the same factual base, the same requested style, the same requested tone, and the same length band. They differ in opening hook, structural emphasis, and which qualifications they foreground, NOT in tone/style/length.

STYLE (apply uniformly across all 3 variants):
- formal: traditional business language, formal salutation ("Dear Hiring Manager,"), conservative paragraph structure, no contractions.
- professional: polished and modern, confident but warm, suits most corporate roles. Contractions OK.
- job_matched: read the tone of the JOB description and mirror it. Casual startup posting: conversational and direct. Bank/law/government: formal. Engineering-heavy: precise and concrete.
- Punctuation: never use em-dashes (the long dash character) or en-dashes in the letter text. They read as AI-generated. Use commas, colons, periods, or parentheses instead.

TONE (apply uniformly across all 3 variants):
- confident: assertive, claims expertise plainly, leads with achievements ("I led", "I delivered", "I built"). No hedging language.
- humble: modest and appreciative, emphasises eagerness to learn and contribute, frames achievements as collaborative ("our team", "I had the opportunity to").
- enthusiastic: energetic, conveys genuine excitement about the role/company, uses vivid verbs and warmer adjectives. Still professional, never gushing.

LENGTH (apply to each variant; word count means visible body words, including greeting and signature):
- short: 200 to 300 words (target ~250)
- medium: 300 to 400 words (target ~350)
- long: 400 to 500 words (target ~450)
Each variant should land within ±15% of the target. Do not pad with filler.

ADDITIONAL CONTEXT:
- If the user message includes an "ADDITIONAL CONTEXT" section, weave it into each variant naturally (e.g. relocation availability, target start date, specific motivation). Do not parrot it verbatim. If absent, ignore it and do not invent context.

OUTPUT FORMAT: return ONLY a valid JSON object with this exact structure. No markdown fences, no preamble, no commentary:

{
  "variants": [
    {"content": "<full variant 1 text, achievement-first>"},
    {"content": "<full variant 2 text, motivation-first>"},
    {"content": "<full variant 3 text, company-alignment-first>"}
  ]
}

The `variants` array MUST have exactly 3 elements. Each `content` MUST be a non-empty string containing the full letter.
"""


class OpenAIService(AIServiceInterface):
    def __init__(self) -> None:
        self._client = AsyncOpenAI(
            api_key=settings.openai_api_key,
            timeout=_TIMEOUT,
            max_retries=0,  # we handle retries/errors ourselves
        )

    async def parse_job_description(self, text: str) -> ParseJobResult:
        empty_data = ParsedJobData()

        try:
            start_ms = time.monotonic()

            response = await self._client.chat.completions.create(
                model=_MODEL,
                response_format={"type": "json_object"},
                messages=[
                    {"role": "system", "content": _JOB_PARSE_SYSTEM_PROMPT},
                    {"role": "user", "content": text},
                ],
                temperature=0,
            )

            elapsed_ms = int((time.monotonic() - start_ms) * 1000)

            raw_content = response.choices[0].message.content or ""
            parsed = self._parse_job_response(raw_content)

            usage_info = AIUsageInfo(
                model=response.model,
                tokens_input=response.usage.prompt_tokens if response.usage else 0,
                tokens_output=response.usage.completion_tokens if response.usage else 0,
                response_time_ms=elapsed_ms,
            )

            return ParseJobResult(data=parsed, usage=usage_info, success=True)

        except Exception as exc:
            logger.warning("OpenAI parse_job_description failed: %s", exc)
            return ParseJobResult(data=empty_data, usage=None, success=False)

    async def generate_cover_letter(
        self,
        source_data: dict,
        job_data: dict,
        preferences: dict,
    ) -> GenerateCoverLetterResult:
        """Generate exactly 3 distinct cover letter variants in a single OpenAI call.

        See `AIServiceInterface.generate_cover_letter` for the contract. The
        implementation mirrors `generate_tailored_resume`: structured JSON
        output via `response_format={"type": "json_object"}`, post-validated
        to enforce exactly 3 non-empty variants. On any failure we return
        `success=False` with an empty variants list so the caller can map to
        a 502 / refund the user's CL credit.

        Failure modes (all return success=False, no exception raised to caller):
        - OpenAI / network / timeout error
        - Invalid JSON returned by the model
        - `variants` list missing, wrong type, or != 3 elements
        - Any variant has empty `content`
        """
        empty_result = GenerateCoverLetterResult(variants=[], usage=None, success=False)

        try:
            # --- Normalise inputs -------------------------------------------------
            style = str(preferences.get("style") or "job_matched")
            tone = str(preferences.get("tone") or "confident")
            length = str(preferences.get("length") or "medium")
            source_type = str(preferences.get("source_type") or "profile")
            additional_context = str(preferences.get("additional_context") or "").strip()

            # Length band → midpoint target word count (used in the user prompt
            # so the model has an explicit target rather than the band only).
            band = _CL_LENGTH_BANDS.get(length, _CL_LENGTH_BANDS["medium"])
            target_words = (band[0] + band[1]) // 2

            # --- Build the user payload ------------------------------------------
            # Keep contact PII out of the OpenAI payload: blank everything in
            # personal_info except full_name. No restore needed — the output is
            # plain letter text. Use a shallow top-level copy + the fresh inner
            # dict the helper returns so the caller's source_data is never mutated.
            sanitized_source = dict(source_data or {})
            sanitized_source["personal_info"] = _strip_contact_pii(
                sanitized_source.get("personal_info")
            )

            user_payload: dict[str, Any] = {
                "source": {
                    # source_type is informational only — the AI must NOT branch
                    # behaviour on it (per system prompt). Surfaced so the model
                    # can reason about provenance if it wants to.
                    "source_type": source_type,
                    "data": sanitized_source,
                },
                "job": {
                    "job_title": job_data.get("job_title") or "",
                    "company": job_data.get("company") or "",
                    "requirements": list(job_data.get("requirements") or []),
                    "description": job_data.get("description") or "",
                    # Pass-through any other helpful fields (location, salary, ...)
                    **{
                        k: v for k, v in (job_data or {}).items()
                        if k not in {"job_title", "company", "requirements", "description"}
                    },
                },
                "preferences": {
                    "style": style,
                    "tone": tone,
                    "length": length,
                    "target_word_count": target_words,
                    "word_count_band": {"min": band[0], "max": band[1]},
                },
            }
            if additional_context:
                # Inserted verbatim per CL-1 spec — only when non-empty.
                user_payload["additional_context"] = additional_context

            user_content = json.dumps(user_payload, ensure_ascii=False)

            # --- AI call ---------------------------------------------------------
            model_name = settings.ai_model_cover_letter or _MODEL

            start_ms = time.monotonic()

            response = await self._client.chat.completions.create(
                model=model_name,
                response_format={"type": "json_object"},
                timeout=_COVER_LETTER_TIMEOUT,
                messages=[
                    {"role": "system", "content": _COVER_LETTER_SYSTEM_PROMPT},
                    {"role": "user", "content": user_content},
                ],
                # Higher temperature than tailoring (0.2): variant distinctness
                # benefits from more creative variation. Still well below the
                # default (1.0) to keep tone/style/length consistent across
                # variants.
                temperature=0.7,
            )

            elapsed_ms = int((time.monotonic() - start_ms) * 1000)

            raw_content = response.choices[0].message.content or ""
            variants = self._parse_cover_letter_response(raw_content)

            usage_info = AIUsageInfo(
                model=response.model,
                tokens_input=response.usage.prompt_tokens if response.usage else 0,
                tokens_output=response.usage.completion_tokens if response.usage else 0,
                response_time_ms=elapsed_ms,
            )

            # Strict post-validation: exactly 3 non-empty variants. Anything
            # else is a hard failure — caller refunds the credit upstream.
            if len(variants) != _CL_VARIANT_COUNT:
                logger.warning(
                    "OpenAI generate_cover_letter expected %d variants, got %d "
                    "(tokens_in=%d tokens_out=%d elapsed_ms=%d)",
                    _CL_VARIANT_COUNT,
                    len(variants),
                    usage_info.tokens_input,
                    usage_info.tokens_output,
                    elapsed_ms,
                )
                return empty_result

            if any(not v.content.strip() for v in variants):
                logger.warning(
                    "OpenAI generate_cover_letter returned an empty variant "
                    "(tokens_in=%d tokens_out=%d elapsed_ms=%d)",
                    usage_info.tokens_input,
                    usage_info.tokens_output,
                    elapsed_ms,
                )
                return empty_result

            # Word-count drift is logged-only — surface to the user rather
            # than failing the batch.
            for i, v in enumerate(variants, start=1):
                wc = len(v.content.split())
                if wc < band[0] * 0.85 or wc > band[1] * 1.15:
                    logger.info(
                        "CL variant %d word_count=%d outside ±15%% of band %s "
                        "(target=%d) — surfacing anyway",
                        i, wc, band, target_words,
                    )

            logger.info(
                "Cover letter generation complete: variants=%d style=%s tone=%s "
                "length=%s source_type=%s tokens_in=%d tokens_out=%d elapsed_ms=%d",
                len(variants),
                style,
                tone,
                length,
                source_type,
                usage_info.tokens_input,
                usage_info.tokens_output,
                elapsed_ms,
            )

            return GenerateCoverLetterResult(
                variants=variants, usage=usage_info, success=True
            )

        except Exception as exc:
            logger.warning("OpenAI generate_cover_letter failed: %s", exc)
            return empty_result

    async def parse_resume_to_profile(
        self,
        file_bytes: bytes,
        file_format: str,
    ) -> ParseResumeProfileResult:
        """Extract a v2.1 ProfileData JSONB structure from an uploaded resume.

        File bytes are processed entirely in-memory (BytesIO) — no temp files
        are written to disk (PRD 3.3.16).

        Raises:
            ValueError: Unsupported format, corrupted file, or text too short
                to be a real resume. The router maps this to HTTP 422.
            Exception: OpenAI/network/validation failures are re-raised
                untouched so the router can map them to HTTP 502.
        """
        # Step A: Extract text in-memory. ValueError here means bad input → 422.
        try:
            text = self._extract_resume_text(file_bytes, file_format)
        except ValueError:
            raise
        except Exception as exc:
            # pdfplumber / python-docx parse failure → corrupted file → 422
            logger.warning("Resume text extraction failed: %s", exc)
            raise ValueError(f"Could not read {file_format} file: {exc}") from exc

        if not text or len(text.strip()) < 50:
            raise ValueError("File contains no extractable text or is too short")

        # Step B: Call OpenAI with structured output. Exceptions re-raised → 502.
        start_ms = time.monotonic()

        response = await self._client.chat.completions.create(
            model=settings.ai_model_parse_resume,
            response_format={"type": "json_object"},
            timeout=_PROFILE_PARSE_TIMEOUT,
            messages=[
                {"role": "system", "content": _PROFILE_PARSE_SYSTEM_PROMPT},
                {"role": "user", "content": text},
            ],
            temperature=0,
        )

        elapsed_ms = int((time.monotonic() - start_ms) * 1000)

        raw_content = response.choices[0].message.content or ""
        profile_dict = self._parse_profile_response(raw_content)

        # Validate against ProfileData (raises pydantic.ValidationError on schema
        # mismatch). Imported lazily so the AI module stays import-light.
        from app.schemas.profile import ProfileData

        validated = ProfileData.model_validate(profile_dict)
        validated_dict = validated.model_dump(mode="json")

        usage_info = AIUsageInfo(
            model=response.model,
            tokens_input=response.usage.prompt_tokens if response.usage else 0,
            tokens_output=response.usage.completion_tokens if response.usage else 0,
            response_time_ms=elapsed_ms,
        )

        logger.info(
            "Resume → ProfileData parse complete: tokens_in=%d tokens_out=%d elapsed_ms=%d",
            usage_info.tokens_input,
            usage_info.tokens_output,
            elapsed_ms,
        )

        return ParseResumeProfileResult(
            profile_data=validated_dict,
            usage=usage_info,
            success=True,
        )

    async def generate_tailored_resume(
        self,
        profile_data: dict,
        job_data: dict,
        preferences: dict,
    ) -> GenerateTailoredResumeResult:
        """Generate a tailored resume from the user's profile and a parsed job.

        See `AIServiceInterface.generate_tailored_resume` for the contract.

        Failure modes (all return success=False, no exception raised):
        - OpenAI/network/timeout error
        - Invalid JSON returned by the model
        - Pydantic validation failure on the result shape

        On success, the returned dict-shaped fields have already been validated
        through `TailoredResumeGenerationResult` so the caller can persist them
        without re-validation.
        """
        empty_result = GenerateTailoredResumeResult(
            tailored_data=_empty_profile_data(),
            matched_keywords=[],
            applied_changes={},
            match_score=0,
            usage=None,
            success=False,
        )

        try:
            # Keep contact PII out of the OpenAI payload: blank everything in
            # personal_info except full_name. We restore the real contacts onto
            # the structured response after the call. Use a shallow top-level
            # copy + the fresh inner dict the helper returns so the caller's
            # profile_data is never mutated.
            original_personal_info = (profile_data or {}).get("personal_info")
            sanitized_profile = dict(profile_data or {})
            sanitized_profile["personal_info"] = _strip_contact_pii(original_personal_info)

            user_payload = {
                "profile": sanitized_profile,
                "job": {
                    "job_title": job_data.get("job_title") or "",
                    "company": job_data.get("company") or "",
                    "requirements": list(job_data.get("requirements") or []),
                    "description": job_data.get("description") or "",
                    # Pass-through other helpful fields if present (location, salary, etc.)
                    **{
                        k: v for k, v in job_data.items()
                        if k not in {"job_title", "company", "requirements", "description"}
                    },
                },
                "preferences": {
                    "sections_order": list(preferences.get("sections_order") or []),
                    "font": preferences.get("font") or "",
                    "template": preferences.get("template") or "",
                },
            }

            user_content = json.dumps(user_payload, ensure_ascii=False)

            start_ms = time.monotonic()

            response = await self._client.chat.completions.create(
                model=settings.ai_model_tailor_resume,
                response_format={"type": "json_object"},
                timeout=_TAILOR_TIMEOUT,
                messages=[
                    {"role": "system", "content": _TAILOR_RESUME_SYSTEM_PROMPT},
                    {"role": "user", "content": user_content},
                ],
                temperature=0.2,
            )

            elapsed_ms = int((time.monotonic() - start_ms) * 1000)

            raw_content = response.choices[0].message.content or ""

            try:
                raw_obj: dict[str, Any] = json.loads(raw_content)
            except json.JSONDecodeError as exc:
                logger.warning("OpenAI generate_tailored_resume returned invalid JSON: %s", exc)
                return empty_result
            if not isinstance(raw_obj, dict):
                logger.warning("OpenAI generate_tailored_resume response was not a JSON object")
                return empty_result

            # Post-process: clamp match_score, normalise color_ids, drop empty
            # applied_changes sections, and ensure tailored_data has all keys.
            tailored_dict = self._coerce_tailored_data(
                raw_obj.get("tailored_data"), profile_data
            )
            # Restore the real contact details (blanked before the AI call for
            # privacy) onto the structured response before validation/persist.
            tailored_dict["personal_info"] = original_personal_info
            matched_keywords = self._coerce_matched_keywords(raw_obj.get("matched_keywords"))
            applied_changes = self._coerce_applied_changes(raw_obj.get("applied_changes"))
            match_score = self._clamp_match_score(raw_obj.get("match_score"))

            usage_info = AIUsageInfo(
                model=response.model,
                tokens_input=response.usage.prompt_tokens if response.usage else 0,
                tokens_output=response.usage.completion_tokens if response.usage else 0,
                response_time_ms=elapsed_ms,
            )

            # Validate the structurally clean dict through the Pydantic schema
            # so the persistence layer (TAILOR-2) can trust the output shape.
            from app.schemas.tailored_resume import TailoredResumeGenerationResult

            try:
                validated = TailoredResumeGenerationResult.model_validate({
                    "tailored_data": tailored_dict,
                    "matched_keywords": matched_keywords,
                    "applied_changes": applied_changes,
                    "match_score": match_score,
                })
            except Exception as exc:  # pydantic.ValidationError or anything else
                logger.warning(
                    "TailoredResumeGenerationResult validation failed: %s", exc
                )
                # Return the partial data so the caller can log/diagnose.
                return GenerateTailoredResumeResult(
                    tailored_data=tailored_dict,
                    matched_keywords=matched_keywords,
                    applied_changes=applied_changes,
                    match_score=match_score,
                    usage=usage_info,
                    success=False,
                )

            # Round-trip through model_dump so we hand back a JSON-clean dict
            # (UUIDs as str, no datetime weirdness) — matching what TAILOR-2
            # will write to JSONB.
            tailored_clean = validated.tailored_data.model_dump(mode="json")
            keywords_clean = [kw.model_dump(mode="json") for kw in validated.matched_keywords]
            applied_clean = validated.applied_changes.model_dump(mode="json")

            logger.info(
                "Tailored resume generation complete: score=%d keywords=%d sections_changed=%d "
                "tokens_in=%d tokens_out=%d elapsed_ms=%d",
                validated.match_score,
                len(keywords_clean),
                len(applied_clean),
                usage_info.tokens_input,
                usage_info.tokens_output,
                elapsed_ms,
            )

            return GenerateTailoredResumeResult(
                tailored_data=tailored_clean,
                matched_keywords=keywords_clean,
                applied_changes=applied_clean,
                match_score=validated.match_score,
                usage=usage_info,
                success=True,
            )

        except Exception as exc:
            logger.warning("OpenAI generate_tailored_resume failed: %s", exc)
            return empty_result

    # ------------------------------------------------------------------
    # Private helpers
    # ------------------------------------------------------------------

    def _parse_job_response(self, content: str) -> ParsedJobData:
        """Parse JSON response from OpenAI into ParsedJobData. Returns empty data on any error."""
        try:
            raw: dict[str, Any] = json.loads(content)
        except json.JSONDecodeError as exc:
            logger.warning("Failed to decode OpenAI JSON response: %s", exc)
            return ParsedJobData()

        requirements_raw = raw.get("requirements")
        requirements: list[str] = []
        if isinstance(requirements_raw, list):
            requirements = [str(r) for r in requirements_raw if r]

        def _str_or_none(val: Any) -> str | None:
            if val is None or val == "":
                return None
            return str(val)

        return ParsedJobData(
            job_title=_str_or_none(raw.get("job_title")),
            company=_str_or_none(raw.get("company")),
            requirements=requirements,
            location=_str_or_none(raw.get("location")),
            salary=_str_or_none(raw.get("salary")),
        )

    def _parse_cover_letter_response(self, content: str) -> list[CoverLetterVariant]:
        """Parse the OpenAI JSON response into a list of CoverLetterVariant.

        Returns the variants list as-extracted (no truncation, no padding).
        The caller is responsible for enforcing the exactly-3 contract — this
        helper just returns what the model gave us, dropping malformed entries.

        Returns [] on JSON decode failure, missing/non-list `variants` field,
        or if every entry is malformed.
        """
        try:
            raw: dict[str, Any] = json.loads(content)
        except json.JSONDecodeError as exc:
            logger.warning("Failed to decode cover letter JSON response: %s", exc)
            return []

        if not isinstance(raw, dict):
            logger.warning("Cover letter response was not a JSON object")
            return []

        variants_raw = raw.get("variants")
        if not isinstance(variants_raw, list):
            logger.warning("Cover letter response missing 'variants' list")
            return []

        variants: list[CoverLetterVariant] = []
        for item in variants_raw:
            # Tolerate models that hand back a bare string instead of {content: ...}.
            if isinstance(item, str):
                content_text = item.strip()
            elif isinstance(item, dict):
                content_text = str(item.get("content", "")).strip()
            else:
                continue
            if content_text:
                variants.append(CoverLetterVariant(content=content_text))

        return variants

    # ------------------------------------------------------------------
    # Profile parsing helpers (PROFILE-2)
    # ------------------------------------------------------------------

    def _extract_resume_text(self, file_bytes: bytes, file_format: str) -> str:
        """Extract plain text from PDF or DOCX bytes. NO temp files written.

        Raises ValueError for unsupported formats. Lower-level library errors
        propagate (the caller wraps them into ValueError).
        """
        from io import BytesIO

        fmt = (file_format or "").lower().lstrip(".")

        if fmt == "pdf":
            import pdfplumber

            text_parts: list[str] = []
            detected_urls: list[str] = []
            with pdfplumber.open(BytesIO(file_bytes)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    if page_text:
                        text_parts.append(page_text)
                    # Collect any hyperlink URIs embedded in the PDF.
                    try:
                        hyperlinks = page.hyperlinks or []
                        for link in hyperlinks:
                            uri = link.get("uri") if isinstance(link, dict) else None
                            if uri and isinstance(uri, str):
                                detected_urls.append(uri.strip())
                        if not hyperlinks:
                            # Fall back to annotations if hyperlinks list is empty.
                            for annot in (page.annots or []):
                                if not isinstance(annot, dict):
                                    continue
                                uri = annot.get("uri")
                                if not uri:
                                    data = annot.get("data") or {}
                                    if isinstance(data, dict):
                                        a = data.get("A") or {}
                                        if isinstance(a, dict):
                                            uri = a.get("URI")
                                if uri and isinstance(uri, (str, bytes)):
                                    if isinstance(uri, bytes):
                                        try:
                                            uri = uri.decode("utf-8", "ignore")
                                        except Exception:
                                            uri = ""
                                    if uri:
                                        detected_urls.append(uri.strip())
                    except Exception:
                        pass  # be defensive — never let hyperlink extraction break text extraction
            body = "\n\n".join(text_parts)
            if detected_urls:
                unique_urls = list(dict.fromkeys(detected_urls))  # dedupe, keep order
                body += "\n\n[DETECTED_URLS]\n" + "\n".join(unique_urls)
            return body

        if fmt == "docx":
            from docx import Document

            doc = Document(BytesIO(file_bytes))
            text_parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            text_parts.append(cell_text)
            # Collect hyperlink URIs from the DOCX relationships.
            detected_urls: list[str] = []
            try:
                for rel in doc.part.rels.values():
                    reltype = getattr(rel, "reltype", "") or ""
                    if "hyperlink" in reltype.lower():
                        target = getattr(rel, "target_ref", "") or ""
                        if target and isinstance(target, str):
                            detected_urls.append(target.strip())
            except Exception:
                pass  # be defensive — never let hyperlink extraction break text extraction
            body = "\n".join(text_parts)
            if detected_urls:
                unique_urls = list(dict.fromkeys(detected_urls))
                body += "\n\n[DETECTED_URLS]\n" + "\n".join(unique_urls)
            return body

        raise ValueError(f"Unsupported file format: {file_format!r} (expected 'pdf' or 'docx')")

    def _parse_profile_response(self, content: str) -> dict[str, Any]:
        """Parse the OpenAI JSON response into a ProfileData-shaped dict.

        Tolerates type oddities (defensive coercion). Strips entries that have
        no meaningful content. Raises ValueError on JSON decode failure so the
        caller can surface a clean error.
        """
        try:
            raw: dict[str, Any] = json.loads(content)
        except json.JSONDecodeError as exc:
            raise ValueError(f"AI returned invalid JSON: {exc}") from exc

        if not isinstance(raw, dict):
            raise ValueError("AI response was not a JSON object")

        return {
            "personal_info": self._coerce_personal_info(raw.get("personal_info")),
            "summary": _str_default(raw.get("summary"), ""),
            "employment": self._coerce_employment_list(raw.get("employment")),
            "education": self._coerce_education_list(raw.get("education")),
            "skills": self._coerce_named_list(raw.get("skills")),
            "projects": self._coerce_project_list(raw.get("projects")),
            "languages": self._coerce_named_list(raw.get("languages")),
            "certificates": self._coerce_certificate_list(raw.get("certificates")),
            "achievements": self._coerce_achievement_list(raw.get("achievements")),
            "volunteer": self._coerce_volunteer_list(raw.get("volunteer")),
        }

    def _coerce_personal_info(self, raw: Any) -> dict[str, Any]:
        if not isinstance(raw, dict):
            raw = {}
        social_links_raw = raw.get("social_links") or []
        social_links: list[dict[str, str]] = []
        if isinstance(social_links_raw, list):
            for item in social_links_raw:
                if not isinstance(item, dict):
                    continue
                platform = _str_default(item.get("platform"), "").strip()
                url = _str_default(item.get("url"), "").strip()
                if not platform or not url:
                    continue
                # Reject pseudo-URLs the model occasionally emits (e.g. url == platform name)
                if url.casefold() == platform.casefold():
                    continue
                if len(url) < 5:
                    continue
                if "." not in url and "://" not in url:
                    continue
                # Reject bare root-domain URLs for platforms that require a profile path.
                # e.g. "https://www.linkedin.com" with no /in/<username> is not a profile.
                _normalised = url.lower().split("://", 1)[-1]  # strip scheme
                _normalised = _normalised.rstrip("/")
                if _normalised in _PLATFORM_ROOT_HOSTS:
                    continue
                social_links.append({"platform": platform, "url": url})
        return {
            "full_name": _str_default(raw.get("full_name"), ""),
            "email": _str_default(raw.get("email"), ""),
            "phone": _str_default(raw.get("phone"), ""),
            "location": _nullable_str(raw.get("location")),
            "social_links": social_links,
        }

    def _coerce_employment_list(self, raw: Any) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        if not isinstance(raw, list):
            return items
        for item in raw:
            if not isinstance(item, dict):
                continue
            role = _str_default(item.get("role"), "").strip()
            company = _str_default(item.get("company"), "").strip()
            if not role and not company:
                # Skip empty entries.
                continue
            items.append({
                "role": role,
                "company": company,
                "location": _nullable_str(item.get("location")),
                "start_date": _str_default(item.get("start_date"), ""),
                "end_date": _nullable_str(item.get("end_date")),
                "is_current": bool(item.get("is_current", False)),
                "details": _coerce_str_list(item.get("details")),
            })
        return items

    def _coerce_education_list(self, raw: Any) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        if not isinstance(raw, list):
            return items
        for item in raw:
            if not isinstance(item, dict):
                continue
            degree = _str_default(item.get("degree"), "").strip()
            institution = _str_default(item.get("institution"), "").strip()
            if not degree and not institution:
                continue
            items.append({
                "degree": degree,
                "institution": institution,
                "field_of_study": _nullable_str(item.get("field_of_study")),
                "location": _nullable_str(item.get("location")),
                "start_date": _str_default(item.get("start_date"), ""),
                "end_date": _nullable_str(item.get("end_date")),
                "is_current": bool(item.get("is_current", False)),
                "description": _nullable_str(item.get("description")),
            })
        return items

    def _coerce_named_list(self, raw: Any) -> list[dict[str, str]]:
        """For skills/languages — each item is just {'name': ...}."""
        items: list[dict[str, str]] = []
        if not isinstance(raw, list):
            return items
        for item in raw:
            name: str = ""
            if isinstance(item, dict):
                name = _str_default(item.get("name"), "").strip()
            elif isinstance(item, str):
                name = item.strip()
            if name:
                items.append({"name": name})
        return items

    def _coerce_project_list(self, raw: Any) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        if not isinstance(raw, list):
            return items
        for item in raw:
            if not isinstance(item, dict):
                continue
            name = _str_default(item.get("name"), "").strip()
            description = _str_default(item.get("description"), "").strip()
            if not name and not description:
                continue
            items.append({
                "name": name,
                "description": description,
                "role": _nullable_str(item.get("role")),
                "start_date": _nullable_str(item.get("start_date")),
                "end_date": _nullable_str(item.get("end_date")),
                "is_current": bool(item.get("is_current", False)),
                "technologies": _coerce_str_list(item.get("technologies")),
                "url": _nullable_str(item.get("url")),
                "repository_url": _nullable_str(item.get("repository_url")),
                "details": _coerce_str_list(item.get("details")),
            })
        return items

    def _coerce_certificate_list(self, raw: Any) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        if not isinstance(raw, list):
            return items
        for item in raw:
            if not isinstance(item, dict):
                continue
            name = _str_default(item.get("name"), "").strip()
            if not name:
                continue
            items.append({
                "name": name,
                "issuer": _nullable_str(item.get("issuer")),
                "issue_date": _nullable_str(item.get("issue_date")),
                "expiry_date": _nullable_str(item.get("expiry_date")),
                "credential_url": _nullable_str(item.get("credential_url")),
            })
        return items

    def _coerce_achievement_list(self, raw: Any) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        if not isinstance(raw, list):
            return items
        for item in raw:
            if not isinstance(item, dict):
                continue
            title = _str_default(item.get("title"), "").strip()
            if not title:
                continue
            items.append({
                "title": title,
                "description": _nullable_str(item.get("description")),
                "date": _nullable_str(item.get("date")),
                "issuer": _nullable_str(item.get("issuer")),
            })
        return items

    def _coerce_volunteer_list(self, raw: Any) -> list[dict[str, Any]]:
        items: list[dict[str, Any]] = []
        if not isinstance(raw, list):
            return items
        for item in raw:
            if not isinstance(item, dict):
                continue
            role = _str_default(item.get("role"), "").strip()
            organization = _str_default(item.get("organization"), "").strip()
            if not role and not organization:
                continue
            items.append({
                "role": role,
                "organization": organization,
                "location": _nullable_str(item.get("location")),
                "start_date": _str_default(item.get("start_date"), ""),
                "end_date": _nullable_str(item.get("end_date")),
                "is_current": bool(item.get("is_current", False)),
                "details": _coerce_str_list(item.get("details")),
            })
        return items

    # ------------------------------------------------------------------
    # Tailored-resume helpers (TAILOR-1)
    # ------------------------------------------------------------------

    def _coerce_tailored_data(self, raw: Any, fallback_profile: dict) -> dict[str, Any]:
        """Coerce the model's tailored_data into a ProfileData-shaped dict.

        Guarantees every canonical section key is present so downstream
        Pydantic validation never fails on missing required keys. Falls back
        to the original profile section when the model omits or mangles one.
        """
        if not isinstance(raw, dict):
            raw = {}
        fallback = fallback_profile if isinstance(fallback_profile, dict) else {}

        def _section(key: str, default: Any) -> Any:
            if key in raw and raw[key] is not None:
                return raw[key]
            if key in fallback and fallback[key] is not None:
                return fallback[key]
            return default

        # personal_info: prefer the model's output, but fall back to the
        # original profile's personal_info if the model dropped it. The system
        # prompt forbids modifying personal_info, but defensive fallback keeps
        # validation green even if the model misbehaves.
        personal_info = raw.get("personal_info")
        if not isinstance(personal_info, dict):
            personal_info = fallback.get("personal_info")

        summary = raw.get("summary")
        if summary is None:
            summary = fallback.get("summary", "")
        if not isinstance(summary, str):
            summary = str(summary or "")

        return {
            "personal_info": personal_info,
            "summary": summary,
            "employment": _section("employment", []),
            "education": _section("education", []),
            "skills": _section("skills", []),
            "projects": _section("projects", []),
            "languages": _section("languages", []),
            "certificates": _section("certificates", []),
            "achievements": _section("achievements", []),
            "volunteer": _section("volunteer", []),
        }

    def _coerce_matched_keywords(self, raw: Any) -> list[dict[str, Any]]:
        """Normalise matched_keywords: dedupe, cap at 12, cycle color_id 1..8.

        The model is instructed to assign color_ids, but we re-cycle them
        deterministically by index here so the result is stable even if the
        model returns out-of-range or duplicate ids.
        """
        if not isinstance(raw, list):
            return []
        seen_terms: set[str] = set()
        out: list[dict[str, Any]] = []
        for item in raw:
            if not isinstance(item, dict):
                continue
            term = _str_default(item.get("term"), "").strip()
            if not term:
                continue
            key = term.casefold()
            if key in seen_terms:
                continue
            seen_terms.add(key)
            # Always re-derive color_id from the keyword's position so the
            # cycle is deterministic and bounded — never trust the model.
            color_id = (len(out) % _TAILOR_PALETTE_SIZE) + 1
            out.append({"term": term, "color_id": color_id})
            if len(out) >= _TAILOR_MAX_KEYWORDS:
                break
        return out

    def _coerce_applied_changes(self, raw: Any) -> dict[str, list[str]]:
        """Drop empty / non-list values; coerce each entry to a stripped str."""
        if not isinstance(raw, dict):
            return {}
        out: dict[str, list[str]] = {}
        for key, val in raw.items():
            if not isinstance(key, str) or not key.strip():
                continue
            if not isinstance(val, list):
                continue
            cleaned: list[str] = []
            for entry in val:
                if entry is None:
                    continue
                s = entry if isinstance(entry, str) else str(entry)
                s = s.strip()
                if s:
                    cleaned.append(s)
            if cleaned:
                out[key.strip()] = cleaned
        return out

    def _clamp_match_score(self, raw: Any) -> int:
        try:
            score = int(raw)
        except (TypeError, ValueError):
            return 0
        return max(0, min(100, score))


# ---------------------------------------------------------------------------
# Module-level coercion helpers (used by profile parsing)
# ---------------------------------------------------------------------------

def _str_default(val: Any, default: str) -> str:
    if val is None:
        return default
    if isinstance(val, str):
        return val
    return str(val)


def _nullable_str(val: Any) -> Optional[str]:
    if val is None:
        return None
    if isinstance(val, str):
        s = val.strip()
        return s if s else None
    return str(val)


def _coerce_str_list(val: Any) -> list[str]:
    if not isinstance(val, list):
        return []
    out: list[str] = []
    for item in val:
        if item is None:
            continue
        s = item if isinstance(item, str) else str(item)
        s = s.strip()
        if s:
            out.append(s)
    return out


def _strip_contact_pii(personal_info: Any) -> Any:
    """Return a copy of personal_info with contact details blanked, keeping
    only full_name. Non-dict / None passes through unchanged. Keeps contact
    PII out of OpenAI payloads (full_name is retained for the cover-letter
    signature)."""
    if not isinstance(personal_info, dict):
        return personal_info
    return {
        **personal_info,
        "email": "",
        "phone": "",
        "location": None,
        "social_links": [],
    }


def _empty_profile_data() -> dict[str, Any]:
    """An empty ProfileData-shaped dict used as fallback on AI failure (TAILOR-1)."""
    return {
        "personal_info": None,
        "summary": "",
        "employment": [],
        "education": [],
        "skills": [],
        "projects": [],
        "languages": [],
        "certificates": [],
        "achievements": [],
        "volunteer": [],
    }
