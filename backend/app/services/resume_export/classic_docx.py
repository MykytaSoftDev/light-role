"""
Classic resume template — DOCX builder.

Ported from app.utils.docx_export, adapted to use the typed ResumeData
Pydantic model instead of raw dict[str, Any].
"""
from __future__ import annotations

import io

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from app.services.resume_export.types import ResumeData


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def build_classic_docx(data: ResumeData) -> bytes:
    """
    Render a ResumeData instance as a classic-style DOCX and return bytes.

    Sections rendered: personal_info, summary, experience, education,
    skills, languages, certifications.
    """
    doc = Document()
    _configure_margins(doc)

    _render_personal_info(doc, data)

    summary = data.summary
    if summary and summary.strip():
        _render_section_heading(doc, "Summary")
        doc.add_paragraph(summary.strip())

    if data.experience:
        _render_section_heading(doc, "Experience")
        for item in data.experience:
            _render_experience_item(doc, item)

    if data.education:
        _render_section_heading(doc, "Education")
        for item in data.education:
            _render_education_item(doc, item)

    if data.skills:
        _render_section_heading(doc, "Skills")
        doc.add_paragraph(", ".join(s for s in data.skills if s))

    if data.languages:
        _render_section_heading(doc, "Languages")
        doc.add_paragraph(", ".join(lang for lang in data.languages if lang))

    if data.certifications:
        _render_section_heading(doc, "Certifications")
        for cert in data.certifications:
            _render_certification_item(doc, cert)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def generate_classic_docx_from_dict(parsed_data: dict) -> bytes:
    """Convenience wrapper: accepts raw dict, parses to ResumeData, calls builder."""
    data = ResumeData.model_validate(parsed_data)
    return build_classic_docx(data)


# ---------------------------------------------------------------------------
# Section renderers
# ---------------------------------------------------------------------------

def _render_personal_info(doc: Document, data: ResumeData) -> None:
    info = data.personal_info
    name = (info.name or "").strip()
    if name:
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(name)
        run.bold = True
        run.font.size = Pt(20)

    contact_parts: list[str] = []
    for value in (info.email, info.phone, info.location, info.linkedin, info.website):
        stripped = (value or "").strip()
        if stripped:
            contact_parts.append(stripped)

    if contact_parts:
        contact_para = doc.add_paragraph(" | ".join(contact_parts))
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in contact_para.runs:
            run.font.size = Pt(10)


def _render_section_heading(doc: Document, heading: str) -> None:
    para = doc.add_paragraph()
    run = para.add_run(heading.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    _add_paragraph_bottom_border(para)


def _render_experience_item(doc: Document, item) -> None:  # item: ExperienceItem
    title = item.title.strip()
    company = item.company.strip()
    start_date = (item.start_date or "").strip()
    end_date = ("Present" if item.current else (item.end_date or "")).strip()
    description = item.description.strip()
    achievements = item.achievements

    heading_parts = [p for p in (title, company) if p]
    heading_text = " — ".join(heading_parts) if heading_parts else "Untitled Position"

    para = doc.add_paragraph()
    run = para.add_run(heading_text)
    run.bold = True
    run.font.size = Pt(11)

    date_parts = [p for p in (start_date, end_date) if p]
    if date_parts:
        date_para = doc.add_paragraph(" – ".join(date_parts))
        for run in date_para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    if description:
        doc.add_paragraph(description)

    for achievement in achievements:
        if achievement and achievement.strip():
            doc.add_paragraph(achievement.strip(), style="List Bullet")


def _render_education_item(doc: Document, item) -> None:  # item: EducationItem
    institution = item.institution.strip()
    degree = item.degree.strip()
    field = (item.field or "").strip()
    start_date = (item.start_date or "").strip()
    end_date = (item.end_date or "").strip()
    gpa = (item.gpa or "").strip()

    degree_parts = [p for p in (degree, field) if p]
    degree_text = ", ".join(degree_parts) if degree_parts else ""

    heading_parts = [p for p in (institution, degree_text) if p]
    heading_text = " — ".join(heading_parts) if heading_parts else "Untitled"

    para = doc.add_paragraph()
    run = para.add_run(heading_text)
    run.bold = True
    run.font.size = Pt(11)

    date_parts = [p for p in (start_date, end_date) if p]
    meta_parts = []
    if date_parts:
        meta_parts.append(" – ".join(date_parts))
    if gpa:
        meta_parts.append(f"GPA: {gpa}")

    if meta_parts:
        meta_para = doc.add_paragraph(" | ".join(meta_parts))
        for run in meta_para.runs:
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def _render_certification_item(doc: Document, item) -> None:  # item: CertificationItem
    name = item.name.strip()
    issuer = (item.issuer or "").strip()
    date = (item.date or "").strip()

    parts = [p for p in (name, issuer, date) if p]
    if parts:
        doc.add_paragraph(" | ".join(parts), style="List Bullet")


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _configure_margins(doc: Document) -> None:
    """Set 1-inch margins on all sides for the default section."""
    from docx.shared import Inches
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)


def _add_paragraph_bottom_border(para) -> None:  # type: ignore[no-untyped-def]
    """Add a thin bottom border to a paragraph via raw OOXML manipulation."""
    from lxml import etree

    pPr = para._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")    # half-points: 4 = 0.5pt
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
