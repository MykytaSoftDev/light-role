"""
Modern resume template — DOCX builder (RTV-2.2).

Approximates the redesigned PDF layout:
  - 2-column borderless table, left col ~2.0" wide with lavender-gray (#F1F2F7) fill.
  - Name + indigo role subtitle placed at the TOP of the right (main) column on white,
    NOT in the sidebar — this matches the PDF redesign where name sits on the white surface.
  - Left column (sidebar): Contact, Skills (individual lines), Languages, Certifications.
  - Right column (main): Name header, Summary, Experience, Education.
  - Section titles: indigo (#4F46E5), uppercase, bold — NO bottom-border underline.
  - Achievement bullets use the › chevron (U+203A) character in indigo, plain paragraph style.
  - Job entries: title bold + dates right-aligned via tab stop; company on next line in indigo.

DOCX-specific tradeoffs vs PDF:
  - The 3pt vertical indigo accent bar (flush-left sidebar edge) is approximated via a
    left cell border (w:left) with color #4F46E5 and width 24 half-points (~3pt). This
    renders correctly in Word; Google Docs and LibreOffice may render it as a thin line only.
  - Skill pills (pill tags with border-radius) are not feasible in DOCX. Each skill is its
    own paragraph line in the sidebar preceded by a › glyph — closest readable approximation.
  - Date right-alignment on the same line as job title uses a right tab stop at the right
    edge of the column rather than a nested table — simpler and more portable across viewers.
  - Letter-spacing (tracking) on section titles is not supported in DOCX XML; omitted.
  - Font: Inter is a system font on most modern OS installs. We specify it by name and rely
    on system availability; DOCX does not bundle fonts. Falls back to Calibri/default if absent.
"""
from __future__ import annotations

import io
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips
from lxml import etree

from app.services.resume_export.types import ResumeData

# ── Color constants ───────────────────────────────────────────────────────────
_INDIGO = RGBColor(0x4F, 0x46, 0xE5)       # #4F46E5 — section titles, role, accents
_SLATE_800 = RGBColor(0x1E, 0x29, 0x3B)    # #1E293B — body text
_SLATE_900 = RGBColor(0x0F, 0x17, 0x2A)    # #0F172A — name
_SLATE_500 = RGBColor(0x64, 0x74, 0x8B)    # #64748B — dates, meta
_SIDEBAR_BG = "F1F2F7"                      # lavender-gray sidebar fill
_SIDEBAR_ACCENT = "4F46E5"                  # left border accent on sidebar cell

# ── Column widths ─────────────────────────────────────────────────────────────
# Total usable width at 0.75" margins on 8.5" paper = 7.0" — but Word letter = 8.5"
# We target ~180pt sidebar (~2.5") + main (~4.5").  Using Inches as proxy.
_LEFT_COL_WIDTH = Inches(2.5)   # ~180pt — close enough; exact pt mapping varies
_RIGHT_COL_WIDTH = Inches(4.5)

# ── Right-column tab stop for date right-alignment ───────────────────────────
# Tab stop at ~4.3" from the left of the right column (full right edge).
_DATE_TAB_TWIPS = int(4.3 * 1440)  # 1440 twips per inch


def build_modern_docx(data: ResumeData) -> bytes:
    doc = Document()
    _configure_margins(doc)

    # ── 2-column layout table ─────────────────────────────────────────────────
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _remove_table_borders(table)

    left_cell = table.cell(0, 0)
    right_cell = table.cell(0, 1)

    # Enforce column widths
    left_cell.width = _LEFT_COL_WIDTH
    right_cell.width = _RIGHT_COL_WIDTH

    # Lavender-gray sidebar fill
    _set_cell_background(left_cell, _SIDEBAR_BG)

    # 3pt indigo accent bar on left edge of sidebar — approximated via cell left border.
    # Tradeoff: renders as a thin line in all DOCX viewers (Word renders ~3pt accurately;
    # Google Docs and LibreOffice may render thinner). Visually present in all viewers.
    _set_cell_left_border(left_cell, _SIDEBAR_ACCENT, width_half_pts=24)

    # ── Sidebar (left cell) ───────────────────────────────────────────────────
    _clear_cell(left_cell)
    _add_sidebar_padding(left_cell)
    _build_sidebar(left_cell, data)

    # ── Main column (right cell) ──────────────────────────────────────────────
    _clear_cell(right_cell)
    _add_main_padding(right_cell)
    _build_main(right_cell, data)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Layout builders ───────────────────────────────────────────────────────────

def _build_sidebar(cell, data: ResumeData) -> None:
    """Render Contact, Skills, Languages, Certifications into the sidebar cell."""
    info = data.personal_info

    # Contact
    _add_sidebar_heading(cell, "Contact")
    for label, value in [
        ("Email", info.email),
        ("Phone", info.phone),
        ("Location", info.location),
        ("LinkedIn", info.linkedin),
        ("Website", info.website),
    ]:
        if value and value.strip():
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            label_run = p.add_run(f"{label}: ")
            label_run.bold = True
            label_run.font.name = "Inter"
            label_run.font.size = Pt(8)
            label_run.font.color.rgb = _SLATE_500
            value_run = p.add_run(value.strip())
            value_run.font.name = "Inter"
            value_run.font.size = Pt(8)
            value_run.font.color.rgb = _SLATE_800

    # Skills — each skill as its own line with › prefix
    # Tradeoff: pill tags are not achievable in DOCX. Individual lines with a › leading
    # character are the closest readable approximation that preserves per-item clarity.
    if data.skills:
        _add_sidebar_heading(cell, "Skills")
        for skill in data.skills:
            if skill and skill.strip():
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                bullet_run = p.add_run("\u203a ")  # › U+203A single right-pointing angle quotation
                bullet_run.font.name = "Inter"
                bullet_run.font.size = Pt(9)
                bullet_run.font.color.rgb = _INDIGO
                text_run = p.add_run(skill.strip())
                text_run.font.name = "Inter"
                text_run.font.size = Pt(9)
                text_run.font.color.rgb = _SLATE_800

    # Languages
    if data.languages:
        _add_sidebar_heading(cell, "Languages")
        for lang in data.languages:
            if lang and lang.strip():
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(1)
                run = p.add_run(lang.strip())
                run.font.name = "Inter"
                run.font.size = Pt(9)
                run.font.color.rgb = _SLATE_800

    # Certifications
    if data.certifications:
        _add_sidebar_heading(cell, "Certifications")
        for cert in data.certifications:
            parts = [x for x in (cert.name, cert.issuer, cert.date) if x and x.strip()]
            if parts:
                p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.space_after = Pt(2)
                name_run = p.add_run(cert.name.strip() if cert.name else "")
                name_run.bold = True
                name_run.font.name = "Inter"
                name_run.font.size = Pt(8)
                name_run.font.color.rgb = _SLATE_800
                if cert.issuer or cert.date:
                    meta = " · ".join(x for x in (cert.issuer, cert.date) if x and x.strip())
                    mp = cell.add_paragraph()
                    mr = mp.add_run(meta)
                    mr.font.name = "Inter"
                    mr.font.size = Pt(8)
                    mr.font.color.rgb = _SLATE_500


def _build_main(cell, data: ResumeData) -> None:
    """Render Name header, Summary, Experience, Education into the main (right) cell."""
    info = data.personal_info

    # ── Name — large, dark, top of main column ────────────────────────────────
    name_para = cell.add_paragraph()
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run((info.name or "Your Name").strip())
    name_run.bold = True
    name_run.font.name = "Inter"
    name_run.font.size = Pt(22)
    name_run.font.color.rgb = _SLATE_900

    # ── Role subtitle (derived from first experience title) ───────────────────
    # Spec: "if no explicit role, use the most recent experience title."
    # ResumeData has no standalone role field, so we derive from experience[0].title.
    role_text: Optional[str] = None
    if data.experience and data.experience[0].title.strip():
        role_text = data.experience[0].title.strip()
    if role_text:
        role_para = cell.add_paragraph()
        role_para.paragraph_format.space_before = Pt(2)
        role_para.paragraph_format.space_after = Pt(8)
        role_run = role_para.add_run(role_text)
        role_run.font.name = "Inter"
        role_run.font.size = Pt(13)
        role_run.font.color.rgb = _INDIGO
        # weight 500 approximated via non-bold (no native weight 500 in DOCX)

    # ── Summary ───────────────────────────────────────────────────────────────
    if data.summary and data.summary.strip():
        _add_main_heading(cell, "Summary")
        p = cell.add_paragraph(data.summary.strip())
        for run in p.runs:
            run.font.name = "Inter"
            run.font.size = Pt(10)
            run.font.color.rgb = _SLATE_800
        p.paragraph_format.space_after = Pt(6)

    # ── Experience ────────────────────────────────────────────────────────────
    if data.experience:
        _add_main_heading(cell, "Experience")
        for item in data.experience:
            _render_experience(cell, item)

    # ── Education ────────────────────────────────────────────────────────────
    if data.education:
        _add_main_heading(cell, "Education")
        for item in data.education:
            _render_education(cell, item)


# ── Section renderers ─────────────────────────────────────────────────────────

def _render_experience(cell, item) -> None:
    title = (item.title or "").strip()
    company = (item.company or "").strip()
    start_date = (item.start_date or "").strip()
    end_date = ("Present" if item.current else (item.end_date or "")).strip()
    date_str = " \u2013 ".join(x for x in (start_date, end_date) if x)  # en-dash

    # Title + dates on same line via right tab stop
    # Tradeoff: we use a tab stop in a single paragraph rather than a nested table.
    # This is more portable and avoids nested-table corruption in some viewers.
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    _add_right_tab_stop(p, _DATE_TAB_TWIPS)

    title_run = p.add_run(title or "Untitled Position")
    title_run.bold = True
    title_run.font.name = "Inter"
    title_run.font.size = Pt(10)
    title_run.font.color.rgb = _SLATE_800

    if date_str:
        tab_run = p.add_run("\t")
        tab_run.font.size = Pt(9)
        date_run = p.add_run(date_str)
        date_run.font.name = "Inter"
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = _SLATE_500

    # Company line in indigo
    if company:
        cp = cell.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(2)
        company_run = cp.add_run(company)
        company_run.font.name = "Inter"
        company_run.font.size = Pt(10)
        company_run.font.color.rgb = _INDIGO

    # Description (if present)
    if item.description and item.description.strip():
        dp = cell.add_paragraph(item.description.strip())
        for run in dp.runs:
            run.font.name = "Inter"
            run.font.size = Pt(10)
            run.font.color.rgb = _SLATE_800
        dp.paragraph_format.space_after = Pt(2)

    # Achievement bullets — › chevron in indigo, plain paragraph (NOT List Bullet style)
    # Using plain paragraphs avoids the default • glyph and keeps the › character.
    for ach in (item.achievements or []):
        if ach and ach.strip():
            ap = cell.add_paragraph()
            ap.paragraph_format.space_before = Pt(1)
            ap.paragraph_format.space_after = Pt(1)
            ap.paragraph_format.left_indent = Pt(10)
            ap.paragraph_format.first_line_indent = Pt(-10)
            bullet_run = ap.add_run("\u203a ")  # › U+203A
            bullet_run.font.name = "Inter"
            bullet_run.font.size = Pt(9)
            bullet_run.font.color.rgb = _INDIGO
            text_run = ap.add_run(ach.strip())
            text_run.font.name = "Inter"
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = _SLATE_800


def _render_education(cell, item) -> None:
    institution = (item.institution or "").strip()
    degree_parts = [x for x in (item.degree or "", item.field or "") if x.strip()]
    degree_text = ", ".join(degree_parts)
    start_date = (item.start_date or "").strip()
    end_date = (item.end_date or "").strip()
    date_str = " \u2013 ".join(x for x in (start_date, end_date) if x)

    # Institution + dates on same line
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(1)
    _add_right_tab_stop(p, _DATE_TAB_TWIPS)

    inst_run = p.add_run(institution or "Untitled")
    inst_run.bold = True
    inst_run.font.name = "Inter"
    inst_run.font.size = Pt(10)
    inst_run.font.color.rgb = _SLATE_800

    if date_str:
        p.add_run("\t").font.size = Pt(9)
        date_run = p.add_run(date_str)
        date_run.font.name = "Inter"
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = _SLATE_500

    # Degree line
    meta_parts = []
    if degree_text:
        meta_parts.append(degree_text)
    if item.gpa:
        meta_parts.append(f"GPA: {item.gpa}")
    if meta_parts:
        mp = cell.add_paragraph(" \u00b7 ".join(meta_parts))
        mp.paragraph_format.space_before = Pt(0)
        mp.paragraph_format.space_after = Pt(2)
        for run in mp.runs:
            run.font.name = "Inter"
            run.font.size = Pt(9)
            run.font.color.rgb = _SLATE_500


# ── Heading helpers ───────────────────────────────────────────────────────────

def _add_sidebar_heading(cell, text: str) -> None:
    """Indigo uppercase bold section title in sidebar — NO bottom border (new design)."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Inter"
    run.font.size = Pt(8)
    run.font.color.rgb = _INDIGO
    # The old design added a #C7D2FE bottom border here via _add_paragraph_bottom_border.
    # The new design has NO underlines on section titles — removed intentionally.


def _add_main_heading(cell, text: str) -> None:
    """Indigo uppercase bold section title in main column — NO bottom border."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Inter"
    run.font.size = Pt(10)
    run.font.color.rgb = _INDIGO
    # No underline — new design spec removes all section title borders.


# ── Cell / table XML helpers ──────────────────────────────────────────────────

def _configure_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


def _clear_cell(cell) -> None:
    """Remove the default empty paragraph python-docx inserts in new cells."""
    for para in list(cell.paragraphs):
        p = para._element
        p.getparent().remove(p)


def _add_sidebar_padding(cell) -> None:
    """Add a small top spacer paragraph for visual breathing room at the top of the sidebar."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)


def _add_main_padding(cell) -> None:
    """Add top spacer for main column to align with sidebar top padding."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)


def _set_cell_background(cell, hex_color: str) -> None:
    """Set cell fill color via w:shd XML element."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove any existing shd to avoid duplicates
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = etree.SubElement(tcPr, qn("w:shd"))
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _set_cell_left_border(cell, hex_color: str, width_half_pts: int = 24) -> None:
    """
    Set the left border of a table cell to approximate the 3pt indigo accent bar.

    width_half_pts: border width in half-points (24 ≈ 3pt). Word renders this accurately;
    LibreOffice and Google Docs may render thinner but the bar is still visually present.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Get or create tcBorders
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = etree.SubElement(tcPr, qn("w:tcBorders"))
    # Remove existing left border element to avoid duplicates
    for existing in tcBorders.findall(qn("w:left")):
        tcBorders.remove(existing)
    left = etree.SubElement(tcBorders, qn("w:left"))
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(width_half_pts))
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), hex_color)


def _remove_table_borders(table) -> None:
    """Remove all visible borders from a table (outer + inner grid lines)."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = etree.SubElement(tbl, qn("w:tblPr"))
    # Remove existing tblBorders to avoid duplicates
    for existing in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(existing)
    tblBdr = etree.SubElement(tblPr, qn("w:tblBorders"))
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = etree.SubElement(tblBdr, qn(f"w:{side}"))
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")


def _add_right_tab_stop(para, position_twips: int) -> None:
    """
    Add a right-aligned tab stop at position_twips within the paragraph.
    Used to right-align dates on the same line as a job/institution title.
    """
    pPr = para._p.get_or_add_pPr()
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is None:
        tabs_el = etree.SubElement(pPr, qn("w:tabs"))
    tab = etree.SubElement(tabs_el, qn("w:tab"))
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(position_twips))


def _add_paragraph_bottom_border(para, hex_color: str = "AAAAAA") -> None:
    """
    Retained for reference only — NOT used in the new Modern design.
    The new design removes all bottom borders from section titles.
    """
    pPr = para._p.get_or_add_pPr()
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)
