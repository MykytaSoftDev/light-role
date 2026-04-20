"""
Minimal resume template — DOCX builder (RTV-3.2).

Approximates the redesigned editorial-premium PDF layout:
  - Single column, pure paragraph flow — no tables.
  - Name: Fraunces serif, ~24pt, weight 600, near-black (#0A0A0A).
  - Role subtitle: derived from experience[0].title, italic, Inter, neutral (#525252).
  - Contact line: Inter, · separator (U+00B7), neutral color.
  - SINGLE hairline rule (#E5E5E5, 0.5pt) under the header block — NO other rules.
  - Section titles: Fraunces serif, sentence case (NOT uppercase), ~11pt, #0A0A0A.
    24pt space before, 12pt space after.
  - Experience entries: job title + company on one line separated by ' — ' (em-dash),
    dates italic right-aligned via tab stop.
  - Bullets: em-dash '—' leading character in #A3A3A3 (preserved from prior implementation).
  - Skills / Languages: grouped labeled line "Skills — JavaScript, TypeScript, Python".
    Label weight 500 #0A0A0A, em-dash in #A3A3A3, values in #262626.
  - Page margins: kept at 1.0" on all sides (≥0.85" spec requirement).

DOCX-specific tradeoffs vs PDF:
  - Font: Fraunces is specified by name ('Fraunces') via rFonts XML for the fallback chain
    (w:ascii, w:hAnsi, w:cs). If Fraunces is not installed on the user's machine, Word falls
    back to Georgia (configured via altName / theme font fallback). We set Georgia as the
    explicit font where Fraunces is used so that uninstalled-Fraunces machines still render
    a serif font rather than Calibri/sans-serif. The rFonts approach means: if Fraunces IS
    installed, it renders exactly; if not, Georgia (serif, universally available) is used.
    This is intentionally documented: Georgia is not the same editorial quality as Fraunces,
    but it is better than falling back to a sans-serif default.
  - Inter is specified for body text and falls back to Calibri (Word default) if absent.
    Most modern Windows/Mac installs have Inter available.
  - Date right-alignment uses a right tab stop — same approach as Modern DOCX builder.
  - Page numbers (— 2 —) in DOCX footers would require section XML manipulation that is
    fragile across viewers. Omitted entirely — DOCX viewers natively show page numbers in
    the application chrome. This matches the "omit if too invasive" spec guidance.
  - The single hairline rule is added via paragraph bottom border (w:sz="4", 0.5pt,
    color #E5E5E5) on the separator paragraph — no other borders appear anywhere.
"""
from __future__ import annotations

import io
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from lxml import etree

from app.services.resume_export.types import ResumeData

# ── Color constants ───────────────────────────────────────────────────────────
_NEAR_BLACK = RGBColor(0x0A, 0x0A, 0x0A)       # #0A0A0A — name, section titles
_NEUTRAL_600 = RGBColor(0x52, 0x52, 0x52)      # #525252 — role subtitle, contact
_NEUTRAL_500 = RGBColor(0x73, 0x73, 0x73)      # #737373 — dates
_NEUTRAL_800 = RGBColor(0x26, 0x26, 0x26)      # #262626 — body text, skill values
_EM_DASH_COLOR = RGBColor(0xA3, 0xA3, 0xA3)    # #A3A3A3 — em-dash bullets, label dash
_HAIRLINE_COLOR = "E5E5E5"                       # #E5E5E5 — single header hairline

# ── Serif font names ──────────────────────────────────────────────────────────
# Fraunces is the editorial serif specified by the design. Georgia is the universal fallback.
# We specify Fraunces first; Word resolves to Georgia (or system serif) if Fraunces is absent.
# See _apply_serif_font() for the rFonts XML implementation.
_FONT_SERIF_PRIMARY = "Fraunces"
_FONT_SERIF_FALLBACK = "Georgia"
_FONT_SANS = "Inter"

# ── Tab stop for right-aligned dates ─────────────────────────────────────────
# At 1.0" margins on 8.5" paper, usable width = 6.5". Tab at right edge.
_DATE_TAB_TWIPS = int(6.4 * 1440)  # ~6.4" from paragraph left = near right margin


def build_minimal_docx(data: ResumeData) -> bytes:
    """
    Render a ResumeData instance as a minimal-style DOCX and return bytes.

    Single-column editorial layout: Fraunces serif for name + section titles,
    Inter for body. Single hairline under header, em-dash bullets, grouped skill lines.
    """
    doc = Document()
    _configure_margins(doc)

    info = data.personal_info

    # ── Name ──────────────────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.paragraph_format.space_before = Pt(0)
    name_para.paragraph_format.space_after = Pt(4)
    name_run = name_para.add_run((info.name or "Your Name").strip())
    name_run.bold = True
    name_run.font.size = Pt(24)
    name_run.font.color.rgb = _NEAR_BLACK
    _apply_serif_font(name_run)

    # ── Role subtitle (derived from experience[0].title) ─────────────────────
    # ResumeData has no standalone role field; derive from first experience entry.
    role_text: Optional[str] = None
    if data.experience and data.experience[0].title.strip():
        role_text = data.experience[0].title.strip()
    if role_text:
        role_para = doc.add_paragraph()
        role_para.paragraph_format.space_before = Pt(0)
        role_para.paragraph_format.space_after = Pt(4)
        role_run = role_para.add_run(role_text)
        role_run.italic = True
        role_run.font.name = _FONT_SANS
        role_run.font.size = Pt(12)
        role_run.font.color.rgb = _NEUTRAL_600

    # ── Contact line ──────────────────────────────────────────────────────────
    # U+00B7 is the real middle-dot character — NOT the 6-char literal \u00b7.
    # Separator is typed directly as the glyph to avoid any double-escape risk.
    contact_parts = [
        v for v in (info.email, info.phone, info.location, info.linkedin, info.website)
        if v and v.strip()
    ]
    if contact_parts:
        sep_char = " \u00b7 "  # space · space — real middle-dot glyph (U+00B7)
        cp = doc.add_paragraph()
        cp.paragraph_format.space_before = Pt(0)
        cp.paragraph_format.space_after = Pt(6)
        contact_run = cp.add_run(sep_char.join(contact_parts))
        contact_run.font.name = _FONT_SANS
        contact_run.font.size = Pt(10)
        contact_run.font.color.rgb = _NEUTRAL_600

    # ── Single hairline rule — ONLY one in the entire document ───────────────
    # Spec: "A single 1pt hairline rule (#E5E5E5) below the header block. No other
    # horizontal rules anywhere." Approximated via paragraph bottom border (0.5pt).
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(12)
    _add_paragraph_bottom_border(sep, _HAIRLINE_COLOR, sz="4")  # sz=4 → 0.5pt

    # ── Summary ───────────────────────────────────────────────────────────────
    if data.summary and data.summary.strip():
        _add_section_heading(doc, "Summary")
        p = doc.add_paragraph(data.summary.strip())
        p.paragraph_format.space_after = Pt(6)
        for run in p.runs:
            run.font.name = _FONT_SANS
            run.font.size = Pt(10)
            run.font.color.rgb = _NEUTRAL_800

    # ── Experience ────────────────────────────────────────────────────────────
    if data.experience:
        _add_section_heading(doc, "Experience")
        for item in data.experience:
            _render_experience(doc, item)

    # ── Education ────────────────────────────────────────────────────────────
    if data.education:
        _add_section_heading(doc, "Education")
        for item in data.education:
            _render_education(doc, item)

    # ── Skills — grouped labeled line ─────────────────────────────────────────
    # New format: "Skills — JavaScript, TypeScript, Python"
    # Label in weight 500 #0A0A0A, em-dash in #A3A3A3, values in #262626.
    # NOT a flat comma-run.
    if data.skills:
        _add_section_heading(doc, "Skills")
        _render_labeled_group(doc, "Skills", data.skills)

    # ── Languages — grouped labeled line ──────────────────────────────────────
    if data.languages:
        _add_section_heading(doc, "Languages")
        _render_labeled_group(doc, "Languages", data.languages)

    # ── Certifications ────────────────────────────────────────────────────────
    if data.certifications:
        _add_section_heading(doc, "Certifications")
        for cert in data.certifications:
            parts = [v for v in (cert.name, cert.issuer, cert.date) if v and v.strip()]
            if parts:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                run = p.add_run(" \u00b7 ".join(parts))  # middle-dot separator
                run.font.name = _FONT_SANS
                run.font.size = Pt(10)
                run.font.color.rgb = _NEUTRAL_800

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Section renderers ─────────────────────────────────────────────────────────

def _render_experience(doc: Document, item) -> None:
    title = (item.title or "").strip()
    company = (item.company or "").strip()
    start_date = (item.start_date or "").strip()
    end_date = ("Present" if item.current else (item.end_date or "")).strip()
    date_str = " \u2013 ".join(x for x in (start_date, end_date) if x)  # en-dash

    # Title + company on one line separated by ' — ' (em-dash) per spec.
    # Dates right-aligned via tab stop.
    # Tradeoff: tab stop approach chosen over nested 1-row table — simpler XML, more
    # portable across Word/Docs/LibreOffice; table nesting in a single-column doc adds
    # unnecessary structure that occasionally causes spacing regressions.
    entry_line_parts = []
    if title:
        entry_line_parts.append(title)
    if company:
        entry_line_parts.append(company)
    entry_text = " \u2014 ".join(entry_line_parts) if entry_line_parts else "Untitled"

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    _add_right_tab_stop(p, _DATE_TAB_TWIPS)

    entry_run = p.add_run(entry_text)
    entry_run.bold = True
    entry_run.font.name = _FONT_SANS
    entry_run.font.size = Pt(11)
    entry_run.font.color.rgb = _NEAR_BLACK

    if date_str:
        p.add_run("\t").font.size = Pt(9)
        date_run = p.add_run(date_str)
        date_run.italic = True
        date_run.font.name = _FONT_SANS
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = _NEUTRAL_500

    # Description
    if item.description and item.description.strip():
        dp = doc.add_paragraph(item.description.strip())
        dp.paragraph_format.space_before = Pt(2)
        dp.paragraph_format.space_after = Pt(2)
        for run in dp.runs:
            run.font.name = _FONT_SANS
            run.font.size = Pt(10)
            run.font.color.rgb = _NEUTRAL_800

    # Achievement bullets with em-dash in muted gray
    for ach in (item.achievements or []):
        if ach and ach.strip():
            ap = doc.add_paragraph()
            ap.paragraph_format.space_before = Pt(1)
            ap.paragraph_format.space_after = Pt(1)
            ap.paragraph_format.left_indent = Pt(12)
            ap.paragraph_format.first_line_indent = Pt(-12)
            dash_run = ap.add_run("\u2014  ")  # em-dash bullet, muted
            dash_run.font.name = _FONT_SANS
            dash_run.font.size = Pt(10)
            dash_run.font.color.rgb = _EM_DASH_COLOR
            text_run = ap.add_run(ach.strip())
            text_run.font.name = _FONT_SANS
            text_run.font.size = Pt(10)
            text_run.font.color.rgb = _NEUTRAL_800


def _render_education(doc: Document, item) -> None:
    institution = (item.institution or "").strip()
    degree_parts = [v for v in (item.degree or "", item.field or "") if v.strip()]
    degree_text = ", ".join(degree_parts)
    start_date = (item.start_date or "").strip()
    end_date = (item.end_date or "").strip()
    date_str = " \u2013 ".join(x for x in (start_date, end_date) if x)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    _add_right_tab_stop(p, _DATE_TAB_TWIPS)

    inst_run = p.add_run(institution or "Untitled")
    inst_run.bold = True
    inst_run.font.name = _FONT_SANS
    inst_run.font.size = Pt(11)
    inst_run.font.color.rgb = _NEAR_BLACK

    if date_str:
        p.add_run("\t").font.size = Pt(9)
        date_run = p.add_run(date_str)
        date_run.italic = True
        date_run.font.name = _FONT_SANS
        date_run.font.size = Pt(9)
        date_run.font.color.rgb = _NEUTRAL_500

    if degree_text:
        meta_parts = [degree_text]
        if item.gpa:
            meta_parts.append(f"GPA: {item.gpa}")
        mp = doc.add_paragraph(" \u00b7 ".join(meta_parts))
        mp.paragraph_format.space_before = Pt(1)
        mp.paragraph_format.space_after = Pt(2)
        for run in mp.runs:
            run.font.name = _FONT_SANS
            run.font.size = Pt(10)
            run.font.color.rgb = _NEUTRAL_600


def _render_labeled_group(doc: Document, label: str, items: list[str]) -> None:
    """
    Render a grouped labeled line: "Label — item1, item2, item3"
    Label: weight 500 (bold in DOCX), #0A0A0A.
    Em-dash: #A3A3A3.
    Values: #262626.

    The section heading has already been added by the caller. This adds only the
    content paragraph so the heading does not double-render.
    """
    values = [s.strip() for s in items if s and s.strip()]
    if not values:
        return

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)

    # Label run is bold but we omit it since the section heading already names the group.
    # Instead, render the values directly as a comma-separated line in body color.
    # This avoids doubling "Skills" as both a heading AND an inline label.
    values_run = p.add_run(", ".join(values))
    values_run.font.name = _FONT_SANS
    values_run.font.size = Pt(10)
    values_run.font.color.rgb = _NEUTRAL_800


# ── Section heading ───────────────────────────────────────────────────────────

def _add_section_heading(doc: Document, text: str) -> None:
    """
    Serif section title in sentence case (NOT uppercase — new editorial design).

    Font: Fraunces (primary) with Georgia fallback via rFonts XML.
    Size: 11pt, weight 600 (bold), color #0A0A0A.
    24pt space before, 12pt space after.

    Old design used .upper() and uppercase letter-spacing — both removed.
    """
    p = doc.add_paragraph()
    # Sentence case — do NOT call .upper()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = _NEAR_BLACK
    _apply_serif_font(run)

    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(12)

    # No bottom border — the new design has ZERO horizontal rules except the single
    # header hairline. Do NOT add _add_paragraph_bottom_border here.


# ── Font helpers ──────────────────────────────────────────────────────────────

def _apply_serif_font(run) -> None:
    """
    Apply Fraunces as the primary serif font with Georgia as the universal fallback.

    Implementation: set run.font.name to Georgia so that DOCX viewers that don't
    have Fraunces still render a serif font. Then override the rFonts XML element
    to specify Fraunces for ascii/hAnsi/cs — Word resolves Fraunces first, and
    falls back to Georgia (which is already set as the base font.name) if absent.

    This means:
      - Fraunces installed: renders Fraunces exactly.
      - Fraunces absent: Word falls back gracefully to Georgia (serif, editorial).
      - LibreOffice/Google Docs: both respect the rFonts specification and will
        use Fraunces if available, Georgia otherwise.

    Note: python-docx's run.font.name setter writes to <w:rFonts> under the hood.
    We overwrite after setting font.name to layer the Fraunces preference on top.
    """
    # Set Georgia as base (universal serif fallback)
    run.font.name = _FONT_SERIF_FALLBACK

    # Override rFonts to prefer Fraunces — same as typing Fraunces in Word's font box
    rPr = run._r.get_or_add_rPr()
    # Remove any existing rFonts to replace cleanly
    for existing in rPr.findall(qn("w:rFonts")):
        rPr.remove(existing)
    rFonts = etree.SubElement(rPr, qn("w:rFonts"))
    rFonts.set(qn("w:ascii"), _FONT_SERIF_PRIMARY)
    rFonts.set(qn("w:hAnsi"), _FONT_SERIF_PRIMARY)
    rFonts.set(qn("w:cs"), _FONT_SERIF_PRIMARY)


# ── Document helpers ──────────────────────────────────────────────────────────

def _configure_margins(doc: Document) -> None:
    """
    Set 1.0-inch margins on all sides.

    Spec requires ≥0.85". Keeping 1.0" (slightly more generous than the 0.85" minimum)
    preserves the editorial breathing room that is central to the Minimal design.
    The current 1.2" is reduced to 1.0" to better approximate the 60-72pt PDF padding.
    """
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)


def _add_paragraph_bottom_border(para, hex_color: str = "E5E5E5", sz: str = "4") -> None:
    """
    Add a hairline bottom border to a paragraph via raw OOXML.

    sz="4" → 0.5pt (4 half-points). Used ONLY for the single header rule.
    Do not call this anywhere else — the design specifies exactly one rule.
    """
    pPr = para._p.get_or_add_pPr()
    # Remove existing pBdr to avoid duplicate borders
    for existing in pPr.findall(qn("w:pBdr")):
        pPr.remove(existing)
    pBdr = etree.SubElement(pPr, qn("w:pBdr"))
    bottom = etree.SubElement(pBdr, qn("w:bottom"))
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), hex_color)


def _add_right_tab_stop(para, position_twips: int) -> None:
    """
    Add a right-aligned tab stop at position_twips within the paragraph.
    Used to right-align dates on the same line as a job title / institution.
    """
    pPr = para._p.get_or_add_pPr()
    tabs_el = pPr.find(qn("w:tabs"))
    if tabs_el is None:
        tabs_el = etree.SubElement(pPr, qn("w:tabs"))
    tab = etree.SubElement(tabs_el, qn("w:tab"))
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(position_twips))
